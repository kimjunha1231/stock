from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn


ROOT = Path(__file__).resolve().parent
SOURCE = ROOT / "deliverables" / "hyundai_inventory_tech_stack_architecture.docx"
OUTPUT = ROOT / "deliverables" / "hyundai_inventory_tech_stack_architecture_reviewed_v1.1.docx"


def set_cell(cell, text: str) -> None:
    cell.text = text
    for paragraph in cell.paragraphs:
        paragraph.paragraph_format.space_before = Pt(3)
        paragraph.paragraph_format.space_after = Pt(3)
        paragraph.paragraph_format.line_spacing = 1.1
        for run in paragraph.runs:
            set_run_font(run)
            run.font.size = Pt(9)


def set_run_font(run) -> None:
    run.font.name = "Arial Unicode MS"
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        from docx.oxml import OxmlElement
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    rfonts.set(qn("w:ascii"), "Arial Unicode MS")
    rfonts.set(qn("w:hAnsi"), "Arial Unicode MS")
    rfonts.set(qn("w:eastAsia"), "Arial Unicode MS")
    rfonts.set(qn("w:cs"), "Arial Unicode MS")
    lang = rpr.find(qn("w:lang"))
    if lang is None:
        from docx.oxml import OxmlElement
        lang = OxmlElement("w:lang")
        rpr.append(lang)
    lang.set(qn("w:val"), "ko-KR")
    lang.set(qn("w:eastAsia"), "ko-KR")


def replace_in_tables(doc: Document, old: str, new: str) -> int:
    count = 0
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if old in cell.text:
                    set_cell(cell, cell.text.replace(old, new))
                    count += 1
    return count


def replace_in_paragraphs(doc: Document, old: str, new: str) -> int:
    count = 0
    for paragraph in doc.paragraphs:
        if old in paragraph.text:
            for run in paragraph.runs:
                run.text = run.text.replace(old, new)
            count += 1
    return count


def set_row(table, row_index: int, values: list[str]) -> None:
    row = table.rows[row_index]
    for cell, value in zip(row.cells, values):
        set_cell(cell, value)


doc = Document(str(SOURCE))

# Normalize all existing runs as well as new edits so Korean glyphs render
# consistently in Word, LibreOffice, and the document QA renderer.
for paragraph in doc.paragraphs:
    for run in paragraph.runs:
        set_run_font(run)
for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    set_run_font(run)

replace_in_paragraphs(doc, "문서 버전: v1.0", "문서 버전: v1.1")
replace_in_paragraphs(doc, "LLM 구조화 출력(Structured Output) & Redis 0초 캐싱", "LLM 구조화 출력(Structured Output) & 버전 캐시")
replace_in_paragraphs(doc, "동일 재고 조건의 설명 문장은 Redis에 캐싱하여 0초 응답을 보장합니다.", "동일 재고 조건의 설명 문장은 모델·프롬프트·데이터 버전이 포함된 키로 캐싱하며, TTL·무효화·PII 보호를 적용합니다.")
replace_in_paragraphs(doc, "매월 파이썬 ML 모델이 스스로 재학습하여 정밀도를 지속적으로 향상시킵니다.", "시간순 백테스트와 승인·롤백 절차를 통과한 모델만 배포합니다. 검증 게이트 없는 자동 재학습은 실행하지 않습니다.")

# This document is a rebuild target, not a statement about the old repository.
replace_in_tables(doc, "React 19 + Vite", "React 19 + Vite + TypeScript (신규 개발 확정)")
replace_in_tables(doc, "Tailwind CSS v4 + shadcn/ui", "Tailwind CSS v4 + shadcn/ui (조건부 확정)")
replace_in_tables(doc, "Java 21 + Spring Boot 3.3", "Java 21 + Spring Boot 3.x (신규 개발 목표)")
replace_in_tables(doc, "Spring State Machine", "명시적 상태전이 서비스 (MVP; State Machine은 조건부)")
replace_in_tables(doc, "Spring Batch 5 + Quartz", "Spring Batch 5 + 스케줄러 1종 선택")
replace_in_tables(doc, "Spring Security + JWT", "Spring Security OAuth2/OIDC + RBAC")
replace_in_tables(doc, "PostgreSQL 16+", "PostgreSQL 16+ (신규 개발 목표)")
replace_in_tables(doc, "Redis 7.x", "Redis 7.x (선택적 캐시)")
replace_in_tables(doc, "PuLP / SciPy", "OR-Tools/HiGHS 또는 PuLP + 명시적 solver")
replace_in_tables(doc, "scikit-learn / LightGBM", "scikit-learn / LightGBM + 시계열 검증")
replace_in_tables(doc, "Google Gemini 1.5 / Ollama / GPT-4o-mini", "LLM Provider Adapter (OpenAI/Gemini/Ollama 등)")
replace_in_tables(doc, "동시 접속자 500명 부하 상황의 P95 Latency 및 TPS 측정", "실제 피크·SLO에 따른 P95/P99·오류율·처리량 측정")

# Rewrite the rows whose original wording contained unsupported guarantees or
# duplicated responsibilities. Keeping the same table shape preserves the
# source document's traceability while making the decision text executable.
t0, t1, t2, t3, t4 = doc.tables
set_row(t0, 1, [
    "React 19 + Vite + TypeScript (신규 개발 확정)",
    "B2B 오퍼레이션 타워 SPA, 빠른 HMR·빌드, 클라이언트 시뮬레이션",
    "vs Next.js / 내부 운영 도구는 SEO보다 API 분리·배포 단순성·반응성이 중요하여 Vite 선택",
    "SSR·서버 액션이 필요해지면 BFF 또는 프레임워크를 재검토",
])
set_row(t0, 2, [
    "TanStack Query v5 (도입 예정)",
    "재고 API 캐싱, 상세진단 비동기 로딩, 전략생성 폴링",
    "vs fetch 직접 관리 / 캐시·무효화·재시도 정책을 표준화",
    "실 API 계약 확정 후 도입",
])
set_row(t0, 4, [
    "Zustand (도입 후보)",
    "시뮬레이션 UI 상태·다중 선택 상태의 화면 간 공유",
    "vs React state / 여러 화면으로 확장될 때 selector 기반 구독",
    "서버 상태와 UI 상태를 먼저 분리",
])
set_row(t0, 7, [
    "Tailwind CSS v4 + shadcn/ui (조건부 확정)",
    "현대백화점 그린 디자인 토큰·접근성 컴포넌트·상태 배지",
    "vs MUI / 필요한 컴포넌트만 소유하여 스타일 일관성 확보",
    "버전 고정·디자인 토큰·접근성 검토 필요",
])
set_row(t0, 10, [
    "테스트: Vitest + RTL + Playwright (도입 예정)",
    "수식 회귀·컴포넌트·핵심 승인 흐름 E2E 검증",
    "vs Jest/Cypress / Vite·React와의 CI 유지비를 비교해 최소 조합 선택",
    "도구보다 계산·승인 경로의 테스트 우선",
])

set_row(t1, 1, [
    "Java 21 + Spring Boot 3.x (신규 개발 목표)",
    "REST API, 권한·승인·감사·재고 도메인 로직",
    "vs Node.js/Python 단일 백엔드 / 트랜잭션·배치·보안을 한곳에서 관리하고 FastAPI는 계산 서비스로 분리",
    "서비스 간 OpenAPI·timeout·재시도·Trace ID 계약 필요",
])
set_row(t1, 2, [
    "Spring Data JPA + QueryDSL 5.x (목표)",
    "도메인 CRUD·다중 조건 동적 검색",
    "vs Native SQL/jOOQ / 업무 조회는 QueryDSL, 대용량 집계는 별도 read model·SQL",
    "JPA·SQL 매퍼를 처음부터 동시에 넣지 않음",
])
set_row(t1, 4, [
    "명시적 상태전이 서비스 (MVP)",
    "위험탐지→검토→승인→실행→완료의 허용 전이·감사 이력",
    "vs State Machine / 고정 흐름은 Enum·전이 서비스·낙관적 잠금으로 시작",
    "병렬 승인·동적 전이가 늘어날 때만 State Machine 도입",
])
set_row(t1, 5, [
    "Spring Security OAuth2/OIDC + RBAC (목표)",
    "점포 담당자·본사 예외 담당자의 최소권한 분리",
    "vs 직접 JWT 발급 / 사내 IdP 토큰을 Resource Server로 검증",
    "IdP·권한 매핑·감사 로그 필요",
])
set_row(t1, 6, [
    "Spring Batch 5 + 스케줄러 1종 선택",
    "재고 스냅샷·위험도 집계·재시도 가능한 배치",
    "vs @Scheduled/Quartz/Kubernetes CronJob / 단일·다중 인스턴스 요구를 확인해 하나 선택",
    "Batch와 Quartz 동시 사용 금지",
])

set_row(t2, 1, [
    "PostgreSQL 16+ (신규 개발 목표)",
    "재고·판매·원가·승인 이력 정규화, 실행 파라미터·설명 payload만 JSONB",
    "vs MySQL/Oracle / 관계형 조회·감사와 JSONB 유연성 확보",
    "파티션·인덱스·백업·보존기간을 함께 설계",
])
set_row(t2, 2, [
    "Redis 7.x (선택적 캐시)",
    "위험재고 집계·LLM 설명의 짧은 TTL 캐시와 작업 상태 공유",
    "vs 애플리케이션 메모리 / 분산 캐시와 TTL 제공",
    "원본 데이터 아님. 캐시 적중은 0초가 아니며 키 버전·무효화·stampede 방지 필요",
])

set_row(t3, 1, [
    "Python 3.11+ + FastAPI (목표)",
    "수리 최적화·ML 예측 전용 독립 서비스",
    "vs Spring AI 단독 / Python 최적화·ML 생태계를 직접 활용",
    "HTTP 계약·timeout·재시도·CPU 작업 격리 필요",
])
set_row(t3, 2, [
    "OR-Tools/HiGHS 또는 PuLP + 명시적 solver (조건부)",
    "보관·폐기·할인·물류 제약을 포함한 증분 기여현금이익 최적화",
    "vs 직접 작성 / 할인 단계·수량·채널 선택이 정수 제약이면 MILP 사용",
    "‘100% 정확’ 금지. 모델·허용오차·infeasible·solver 로그 기록",
])
set_row(t3, 3, [
    "scikit-learn / LightGBM + 시계열 검증",
    "판매량·소진율·가격 반응 추정",
    "vs RAG/Vector DB / 정형 숫자는 회귀·시계열 모델 대상",
    "시간순 백테스트·내생성·누수·신뢰구간·드리프트·콜드스타트 필요",
])
set_row(t3, 4, [
    "LLM Provider Adapter (조건부)",
    "판단 근거·시나리오 배경·실행 가이드 문장 생성",
    "vs 특정 모델 고정 / 공급자·모델명·가격·한도를 설정으로 교체",
    "PII 마스킹·JSON Schema·주입 방어·토큰 예산·fallback·사람 승인",
])

set_row(t4, 5, [
    "k6 (목표)",
    "승인·검색·시뮬레이션 API의 SLA별 P95/P99·오류율·처리량 측정",
    "vs JMeter/Locust / 실제 피크 사용자와 SLO로 시나리오 산정",
    "동시 사용자 500명은 임의 기준이 아니며 운영 데이터로 재산정",
])
set_row(t4, 1, [
    "Sentry (선택)",
    "프론트·백엔드 런타임 오류와 사용자 영향 추적",
    "vs 중앙 로그만 사용 / 오류 집계와 재현 정보를 빠르게 확인",
    "요금·보존·민감정보 전송 정책을 계약 전에 확인",
])
set_row(t4, 3, [
    "로그 수집기 + Grafana Loki (목표)",
    "구조화 로그의 중앙 수집·검색·보존",
    "vs ELK / 운영 규모와 검색 요구에 따라 수집기 선택",
    "수집기·보존기간·PII 마스킹을 함께 설정",
])

# Append a review section so the original comparison remains traceable while the
# decision rules and caveats are explicit for the rebuild.
doc.add_page_break()
heading = doc.add_paragraph()
heading.style = doc.styles["Heading 1"]
run = heading.add_run("5. 세밀 검토 결과 및 1개월 신규 개발 권고")
set_run_font(run)
run.font.color.rgb = RGBColor(15, 76, 58)

lead = doc.add_paragraph()
lead.paragraph_format.space_after = Pt(8)
lead.add_run(
    "검토 전제: 기존 저장소의 구현 상태는 판단에서 제외하고, 새로 구축할 React 기반 제품의 목표 스택으로 평가했습니다. "
    "기술 이름만으로 정확도·수익·0원 비용이 보장되지는 않으므로 데이터 계약, 모델 검증, 승인 통제를 함께 구현해야 합니다."
)

def bullet(text: str) -> None:
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    set_run_font(run)


bullet("프론트엔드 확정안: React 19 + Vite + TypeScript. 운영 도구는 SEO보다 API 분리·배포 단순성·시뮬레이션 반응성이 중요하므로 Vite 선택이 타당합니다.")
bullet("서버 상태는 TanStack Query, 화면 조율 상태는 Zustand로 분리합니다. 두 라이브러리는 실 API 계약과 상태 경계가 확정된 뒤 도입합니다.")
bullet("Spring Boot와 FastAPI를 분리할 경우 OpenAPI 스키마, timeout·재시도·멱등성·회로차단·Trace ID 전파를 먼저 고정합니다.")
bullet("JPA/QueryDSL과 MyBatis를 처음부터 함께 사용하지 않습니다. CRUD·동적 필터는 QueryDSL, 대용량 집계는 검증된 read model 또는 SQL로 제한합니다.")
bullet("상태가 고정된 MVP는 Enum·전이 서비스·낙관적 잠금·감사 로그로 시작하고, 병렬 승인·동적 전이가 늘어날 때 State Machine을 검토합니다.")
bullet("Spring Batch와 Quartz를 중복 사용하지 않습니다. 배치 재시작·중복 방지 요구를 확인해 Quartz, @Scheduled 또는 외부 스케줄러 중 하나를 선택합니다.")
bullet("PostgreSQL은 재고·판매·원가·승인 이력을 정규화하고 실행 파라미터·설명 payload만 JSONB로 저장합니다. Redis는 원본이 아닌 TTL 캐시이며 캐시 적중을 0초로 표현하지 않습니다.")
bullet("할인 단계·수량·채널 선택이 정수 제약이면 OR-Tools/HiGHS 또는 PuLP와 solver를 사용합니다. 최적화 결과는 입력·모델·허용오차에 의존하므로 infeasible·solver 로그·fallback을 남깁니다.")
bullet("LightGBM 예측은 시간순 백테스트, 할인 내생성·데이터 누수 점검, 신뢰구간·드리프트·콜드스타트 보정을 포함해야 합니다. ‘100% 정확’ 표현은 사용하지 않습니다.")
bullet("LLM은 설명문 생성 보조 역할로 제한하고, Provider Adapter·JSON Schema 검증·PII 마스킹·프롬프트 주입 방어·토큰 예산·사람 승인 조건을 둡니다. 무료 한도·가격은 문서에 고정하지 않습니다.")

heading2 = doc.add_paragraph()
heading2.style = doc.styles["Heading 2"]
run = heading2.add_run("5.1 구축 전에 고정할 운영 계약")
set_run_font(run)
for text in [
    "데이터 계약: SKU·점포·소유권·원가·재고 기준시각, 단위, 결측·중복·지연 처리, 입력 버전.",
    "보안·권한: 사내 IdP/OIDC, 최소권한 RBAC, 승인 전 가격 변경 금지, 모든 추천·승인·실행 감사 로그.",
    "모델 안전: 기준선 대비 증분 기여현금이익, 시간순 백테스트, 신뢰도 임계치, 하드 차단, fallback, 사람 승인.",
    "운영 예산: LLM 토큰·배치·DB·캐시를 건당 측정하고, timeout·재시도·멱등성·백업·복구 목표를 정의.",
]:
    bullet(text)

heading3 = doc.add_paragraph()
heading3.style = doc.styles["Heading 2"]
run = heading3.add_run("5.2 우선순위")
set_run_font(run)
for text in [
    "P0: React/Vite 화면, Spring API 계약, PostgreSQL 스키마, 수식·시뮬레이션 결정 테이블, 승인·감사 로그.",
    "P1: FastAPI 최적화 서비스, 배치 스냅샷, 캐시, 모델 백테스트·모니터링.",
    "P2: LLM 설명 자동화, 고급 State Machine, 다중 채널·다중 점포 확장.",
]:
    bullet(text)

# Footer note for traceability.
note = doc.add_paragraph()
note.paragraph_format.space_before = Pt(10)
note.alignment = WD_ALIGN_PARAGRAPH.RIGHT
run = note.add_run("검토판 v1.1 · 신규 개발 목표 기준 · 2026-07-27")
set_run_font(run)
run.italic = True

doc.save(str(OUTPUT))
print(OUTPUT)
