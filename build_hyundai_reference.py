from __future__ import annotations

from dataclasses import dataclass
from pathlib import Path
from typing import Iterable, Sequence

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = Path("deliverables/hyundai-deptstore-ai-inventory-cost-strategy-reference.docx")
OUT.parent.mkdir(parents=True, exist_ok=True)

# Resolved design preset: compact_reference_guide + editorial_cover header treatment.
# Named override: Noto Sans KR is used consistently; the render QA command supplies a writable font cache.
PAGE_W = 9360
TABLE_INDENT = 120
GREEN = "0F4C3A"
GOLD = "9E7C3B"
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
INK = "17212B"
MUTED = "64748B"
LIGHT = "E8EEF5"
PALE = "F4F6F9"
AMBER = "7A5A00"
RED = "9B1C1C"
GREEN_LIGHT = "EAF5F0"


@dataclass
class Source:
    code: str
    title: str
    url: str
    date: str
    supports: str


SOURCES = [
    Source("ORACLE-RETAIL-INSEASON", "Oracle Retail In-season Planning User Guide", "https://docs.oracle.com/cd/E12444_01/pdf/1301/html_mwa_user/inseason.htm", "접근 2026-07-27", "시즌 중 상품계획(ROS, Sell Thru %, Needed Sales Lift %, Target WOS, WP/CP/OP), 점포 클러스터, 계층적 목표 설정 및 실시간 과잉/품절 경고 체계"),

    Source("HDS-ESG-2024", "현대백화점 지속가능경영보고서 2024 공시", "https://kind.krx.co.kr/external/2025/06/30/000286/20250618000183/61979.htm", "2025-06-30", "보고 범위(백화점 15개·아울렛 9개), GRI/SASB·K-IFRS 기준, 2024년 및 일부 2025년 상반기 성과"),
    Source("HDS-ESG-PAGE", "현대백화점 ESG 평가 정보", "https://ehyundai.com/newPortal/RC/ESG00003_V.do", "접근 2026-07-26", "2025 KCGS 통합 A+, CDP Leadership A, 2024 ISO14001 인증"),
    Source("HDS-CIRCULAR", "현대백화점 자원순환", "https://www.ehyundai.com/newPortal/RC/ENV00003_V.do", "접근 2026-07-26", "폐비닐 열분해 재활용, Project100 재생지·쇼핑백, 365 리사이클, 다회용기·전자영수증"),
    Source("HDS-PROJECT100", "현대백화점그룹 Project 100 소개", "https://www.ehyundai.com/newPortal/group/GN/GN000005.do?seq=1067", "접근 2026-07-26", "점포 폐지 수거→100% 재생지→쇼핑백 순환, 4년간 1,758톤 활용"),
    Source("HDS-FIVEYEAR", "더현대 서울 오픈 5주년 성과", "https://www.ehyundai.com/newPortal/group/GN/GNN000020_V.do?seq=503", "접근 2026-07-26", "2025년 매출 1조2,864억원, 2030 매출비중 58%, 팝업 1,892여 건, 체류시간 37분"),
    Source("HDS-HAI", "현대백화점 생성형 AI 멘토 Hai", "https://www.ehyundai.com/newPortal/group/GN/GNN000020_V.do?seq=434", "접근 2026-07-26", "현대퓨처넷 협업, 13개 직무 139명 전문가 검증, 일평균 350건 이상 질문"),
    Source("HDS-HEYDI", "쇼핑 AI 헤이디와 현대퓨처넷", "https://www.ehyundai.com/newPortal/group/GN/GNN000010_V.do?seq=691", "접근 2026-07-26", "점포 실시간 정보 연동, 브랜드·식당·팝업·전시·프로모션을 대화형 추천에 활용"),
    Source("FUTURENET-JOB", "현대퓨처넷 IT/디지털 직무 소개", "https://recruit.ehyundai.com/job-introduction/job-company-dtl.nhd?coCd=HDHCN&dtlJobCd=017&jobCd=JOB_CD_001", "접근 2026-07-26", "ERP·인프라·보안·메시징·전시·사이니지·운영데이터 분석 역량"),
    Source("HDS-VIP-2025", "현대백화점 2025 VIP 프로그램 변경사항", "https://ehyundai.com/mobile/card/vip/view_2025.do", "2025", "VIP 마일리지·리워드, Club YP 연령 변경, 혜택·적용 제외 조건"),
    Source("HDS-YP-2026", "현대백화점 2026 Club YP 서비스", "https://www.ehyundai.com/HD/VIP/2025/yp/serp.do", "접근 2026-07-26", "VIP 할인 범위 및 임대매장·명품·식품·할인매장·온라인몰 제외 조건"),
    Source("HPOINT", "현대백화점그룹 H.Point", "https://www.h-point.co.kr/stack/charge.nhd", "접근 2026-07-26", "통합 멤버십·포인트 사용 접점"),
    Source("HDS-SUBSCRIPTION", "현대식품관 식품 구독", "https://www.ehyundai.com/newPortal/FS/FS000001_V.do", "접근 2026-07-26", "근거리 배송, 배송일 2일 전 변경, 매장 결제·앱 비대면 결제, 중간 취소 제약"),
    Source("HDS-SERVICE", "현대백화점 서비스 제도", "https://www.ehyundai.com/newPortal/CS/CS000004_V.do", "접근 2026-07-26", "식품 안전 보증·식품 감식관, 냉장 보관, 선물 배송·방문수령 조건"),
    Source("HDS-RETURNS", "현대백화점 교환환불 서비스", "https://www.ehyundai.com/newPortal/CS/CS000003_V.do", "접근 2026-07-26", "7일 교환·환불, 동종 구매조건, 상품군별 반품 예외"),
    Source("HDS-BEEF-2025", "현대백화점 2025 추석 한우 선물세트", "https://www.ehyundai.com/newPortal/group/GN/GNN000010_V.do?seq=795", "2025-08-31", "한우 선물세트 11만 세트, 전년 대비 10% 확대, 산소치환 포장"),
    Source("HDS-POPUP", "더현대 팝업 페스타", "https://www.ehyundai.com/newPortal/NS/NS000001_V.do?bbsCd=201&seq=2090544", "2025", "전점·더현대닷컴 팝업, 한정 기간·한정 물량·이벤트 운영 특성"),
    Source("MFDS-USEBY", "식품 소비기한 표시제 정책브리핑", "https://www.korea.kr/news/healthView.do?newsId=148911057", "접근 2026-07-26", "소비기한과 유통기한의 개념 차이, 보관조건 준수 전제"),
    Source("MFDS-LABEL", "식품 등의 표시·광고에 관한 법률", "https://law.go.kr/LSW/lsInfoP.do?lsiSeq=257727", "접근 2026-07-26", "식품 표시·광고 실증 및 준수 의무"),
    Source("ALLBARO", "올바로시스템 사업장폐기물 안내", "https://www3.allbaro.or.kr/01_wsi/wsi_system_intro.vm", "접근 2026-07-26", "폐기물관리법에 따른 사업장폐기물 전자관리"),
    Source("PIPC-GENAI", "개인정보위 생성형 AI 개인정보 처리 기준", "https://www.korea.kr/news/policyNewsView.do?newsId=148947194", "2025", "생성형 AI 전 과정의 적법성·안전성·내부관리 기준"),
    Source("OPENAI-PRICE", "OpenAI GPT-5 API 가격 안내(예시 스냅샷)", "https://openai.com/index/introducing-gpt-5-for-developers/", "2025-08-07", "모델별 입력·출력 토큰 과금, Batch·캐싱 등 비용 제어 수단의 예시; 배포 전 최신 가격 재확인"),
]


def set_cell_text(cell, text: str, *, bold=False, color=INK, size=9.2, align=None):
    cell.text = ""
    p = cell.paragraphs[0]
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.15
    r = p.add_run(str(text))
    set_run(r, size=size, color=color, bold=bold)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_run(run, *, name="Noto Sans KR", size=11, color=INK, bold=False, italic=False):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "Noto Sans KR")
    run._element.get_or_add_rPr().rFonts.set(qn("w:cs"), "Noto Sans KR")
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    run.bold = bold
    run.italic = italic


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = tcPr.first_child_found_in("w:tcMar")
    if tcMar is None:
        tcMar = OxmlElement("w:tcMar")
        tcPr.append(tcMar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tcMar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tcMar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def shade(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = tcPr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tcPr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_width(cell, width):
    tcPr = cell._tc.get_or_add_tcPr()
    tcW = tcPr.find(qn("w:tcW"))
    if tcW is None:
        tcW = OxmlElement("w:tcW")
        tcPr.append(tcW)
    tcW.set(qn("w:w"), str(width))
    tcW.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths: Sequence[int], *, header_fill=LIGHT, font_size=9.1):
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl = table._tbl
    tblPr = tbl.tblPr
    tblW = tblPr.find(qn("w:tblW"))
    if tblW is None:
        tblW = OxmlElement("w:tblW")
        tblPr.append(tblW)
    tblW.set(qn("w:w"), str(sum(widths)))
    tblW.set(qn("w:type"), "dxa")
    tblInd = tblPr.find(qn("w:tblInd"))
    if tblInd is None:
        tblInd = OxmlElement("w:tblInd")
        tblPr.append(tblInd)
    tblInd.set(qn("w:w"), str(TABLE_INDENT))
    tblInd.set(qn("w:type"), "dxa")
    layout = tblPr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tblPr.append(layout)
    layout.set(qn("w:type"), "fixed")
    grid = tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for r_idx, row in enumerate(table.rows):
        # Keep rows intact in Word/LibreOffice and repeat the header when a long
        # reference table spans pages. This prevents a source row or formula row
        # from being split into an unreadable continuation at the page top.
        trPr = row._tr.get_or_add_trPr()
        cant_split = trPr.find(qn("w:cantSplit"))
        if cant_split is None:
            trPr.append(OxmlElement("w:cantSplit"))
        if r_idx == 0:
            tbl_header = trPr.find(qn("w:tblHeader"))
            if tbl_header is None:
                trPr.append(OxmlElement("w:tblHeader"))
        for c_idx, cell in enumerate(row.cells):
            width = widths[c_idx]
            set_cell_width(cell, width)
            set_cell_margins(cell)
            if r_idx == 0:
                shade(cell, header_fill)
            for p in cell.paragraphs:
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.line_spacing = 1.15
                if r_idx == 0:
                    p.paragraph_format.keep_with_next = True
                for run in p.runs:
                    set_run(run, size=font_size)


def add_hyperlink(paragraph, text, url, color=BLUE):
    part = paragraph.part
    r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    new_run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    color_el = OxmlElement("w:color")
    color_el.set(qn("w:val"), color)
    rPr.append(color_el)
    u = OxmlElement("w:u")
    u.set(qn("w:val"), "single")
    rPr.append(u)
    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:eastAsia"), "Noto Sans KR")
    rFonts.set(qn("w:ascii"), "Noto Sans KR")
    rFonts.set(qn("w:hAnsi"), "Noto Sans KR")
    rFonts.set(qn("w:cs"), "Noto Sans KR")
    rPr.append(rFonts)
    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), "18")
    rPr.append(sz)
    new_run.append(rPr)
    t = OxmlElement("w:t")
    t.text = text
    new_run.append(t)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)
    return hyperlink


def add_page_field(paragraph):
    run = paragraph.add_run()
    fldChar1 = OxmlElement("w:fldChar")
    fldChar1.set(qn("w:fldCharType"), "begin")
    instrText = OxmlElement("w:instrText")
    instrText.set(qn("xml:space"), "preserve")
    instrText.text = " PAGE "
    fldChar2 = OxmlElement("w:fldChar")
    fldChar2.set(qn("w:fldCharType"), "end")
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)
    set_run(run, size=9, color=MUTED)


def configure_styles(doc):
    sec = doc.sections[0]
    sec.page_width = Inches(8.5)
    sec.page_height = Inches(11)
    sec.top_margin = Inches(0.85)
    sec.bottom_margin = Inches(0.8)
    sec.left_margin = Inches(0.9)
    sec.right_margin = Inches(0.9)
    sec.header_distance = Inches(0.35)
    sec.footer_distance = Inches(0.35)
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Noto Sans KR"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Noto Sans KR")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Noto Sans KR")
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Noto Sans KR")
    normal._element.rPr.rFonts.set(qn("w:cs"), "Noto Sans KR")
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25
    for name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, DARK_BLUE, 10, 5),
    ]:
        st = styles[name]
        st.font.name = "Noto Sans KR"
        st._element.rPr.rFonts.set(qn("w:ascii"), "Noto Sans KR")
        st._element.rPr.rFonts.set(qn("w:hAnsi"), "Noto Sans KR")
        st._element.rPr.rFonts.set(qn("w:eastAsia"), "Noto Sans KR")
        st._element.rPr.rFonts.set(qn("w:cs"), "Noto Sans KR")
        st.font.size = Pt(size)
        st.font.bold = True
        st.font.color.rgb = RGBColor.from_string(color)
        st.paragraph_format.space_before = Pt(before)
        st.paragraph_format.space_after = Pt(after)
        st.paragraph_format.keep_with_next = True
        st.paragraph_format.line_spacing = 1.15
    for list_name in ("List Bullet", "List Number"):
        st = styles[list_name]
        st.font.name = "Noto Sans KR"
        st._element.rPr.rFonts.set(qn("w:ascii"), "Noto Sans KR")
        st._element.rPr.rFonts.set(qn("w:hAnsi"), "Noto Sans KR")
        st._element.rPr.rFonts.set(qn("w:eastAsia"), "Noto Sans KR")
        st._element.rPr.rFonts.set(qn("w:cs"), "Noto Sans KR")
        st.font.size = Pt(10.5)
        st.paragraph_format.left_indent = Inches(0.375)
        st.paragraph_format.first_line_indent = Inches(-0.188)
        st.paragraph_format.space_after = Pt(4)
        st.paragraph_format.line_spacing = 1.25


def add_header_footer(doc):
    for section in doc.sections:
        header = section.header
        p = header.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run("현대백화점 · 재고처리 AI 전략 · 비용관리 참고 문서")
        set_run(r, size=8.5, color=MUTED, bold=True)
        footer = section.footer
        fp = footer.paragraphs[0]
        fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        fp.paragraph_format.space_before = Pt(0)
        r = fp.add_run("Internal reference · 2026-07-26 · Page ")
        set_run(r, size=8.5, color=MUTED)
        add_page_field(fp)


def add_para(doc, text: str = "", *, bold_prefix: str | None = None, color=INK, size=10.5, italic=False, align=None, after=6):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.25
    if bold_prefix and text.startswith(bold_prefix):
        r1 = p.add_run(bold_prefix)
        set_run(r1, size=size, color=color, bold=True)
        r2 = p.add_run(text[len(bold_prefix):])
        set_run(r2, size=size, color=color, italic=italic)
    else:
        r = p.add_run(text)
        set_run(r, size=size, color=color, italic=italic)
    return p


def add_rich_para(doc, parts: Iterable[tuple[str, dict]] , *, after=6, align=None):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.25
    for text, opts in parts:
        r = p.add_run(text)
        set_run(r, size=opts.get("size", 10.5), color=opts.get("color", INK), bold=opts.get("bold", False), italic=opts.get("italic", False))
    return p


def add_bullet(doc, text: str, *, level=0, color=INK):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.375 + level * 0.22)
    p.paragraph_format.first_line_indent = Inches(-0.188)
    r = p.add_run(text)
    set_run(r, size=10.3, color=color)
    return p


def add_number(doc, text: str, *, level=0):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.left_indent = Inches(0.375 + level * 0.22)
    p.paragraph_format.first_line_indent = Inches(-0.188)
    r = p.add_run(text)
    set_run(r, size=10.3)
    return p


def add_callout(doc, label: str, text: str, *, fill=PALE, accent=GREEN):
    t = doc.add_table(rows=1, cols=1)
    set_table_geometry(t, [PAGE_W], header_fill=fill, font_size=10)
    cell = t.cell(0, 0)
    shade(cell, fill)
    set_cell_margins(cell, top=140, bottom=140, start=180, end=180)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(label)
    set_run(r, size=10, color=accent, bold=True)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    p2.paragraph_format.line_spacing = 1.2
    r2 = p2.add_run(text)
    set_run(r2, size=10, color=INK)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_table(doc, headers: Sequence[str], rows: Sequence[Sequence[str]], widths: Sequence[int], *, header_fill=LIGHT, font_size=8.9, first_col_bold=False):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for i, h in enumerate(headers):
        set_cell_text(table.rows[0].cells[i], h, bold=True, color=DARK_BLUE, size=font_size)
        shade(table.rows[0].cells[i], header_fill)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], value, bold=first_col_bold and i == 0, size=font_size)
    set_table_geometry(table, widths, header_fill=header_fill, font_size=font_size)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return table


def add_source_ref(doc, code: str, *, prefix="근거: "):
    src = next(s for s in SOURCES if s.code == code)
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(5)
    r = p.add_run(prefix)
    set_run(r, size=8.8, color=MUTED, italic=True)
    r2 = p.add_run(f"[{code}] {src.title}")
    set_run(r2, size=8.8, color=BLUE, italic=True)
    return p


def add_internal_ref(doc, path: str, note: str):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(5)
    r = p.add_run(f"내부 근거: {note} — {path}")
    set_run(r, size=8.8, color=MUTED, italic=True)
    return p


def heading(doc, level: int, text: str):
    p = doc.add_paragraph(text, style=f"Heading {level}")
    for run in p.runs:
        set_run(run, size={1:16, 2:13, 3:12}[level], color={1:BLUE, 2:BLUE, 3:DARK_BLUE}[level], bold=True)
    return p


def cover(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(34)
    p.paragraph_format.space_after = Pt(10)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("FIELD REFERENCE · INVENTORY AI")
    set_run(r, size=10, color=GOLD, bold=True)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run("현대백화점 재고처리 AI 전략\n비용관리·수익최적화 참고 문서")
    set_run(r, size=27, color=GREEN, bold=True)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(20)
    r = p.add_run("현대퓨처넷 협업을 전제로 한 더현대 서울 Anchor 운영·시뮬레이션 설계 자료")
    set_run(r, size=12.5, color=MUTED)
    add_callout(doc, "이 문서의 사용 목적", "현재 프로젝트가 실제 현대백화점 재고를 AI로 처리할 때 무엇을 비용으로 측정하고, 어떤 순서로 재고를 처리해야 최대 이윤과 브랜드·법규 안전을 함께 확보할 수 있는지 판단하기 위한 실무 참고본입니다. 공개 사실, 프로젝트 코드 근거, 제안·추정·검증 필요 항목을 분리했습니다.", fill=GREEN_LIGHT, accent=GREEN)
    meta = [
        ("기준일", "2026-07-26"),
        ("Anchor", "현대백화점 더현대 서울 2F·3F·B1·1F 직매입 우선"),
        ("핵심 사용자", "더현대 서울 재고전략팀 / 계열사 담당자"),
        ("문서 상태", "조사·설계 참고본 — 실운영 전 원천 데이터·계약·법무 검증 필요"),
    ]
    add_table(doc, ["항목", "내용"], meta, [1800, 7560], header_fill="F2F4F7", font_size=9.2, first_col_bold=True)
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(36)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("현대백화점그룹 공개자료 · 프로젝트 내부 문서/코드 · 법규/정책 자료 기반")
    set_run(r, size=9, color=MUTED, italic=True)
    doc.add_page_break()


def build_document():
    doc = Document()
    configure_styles(doc)
    cover(doc)

    heading(doc, 1, "0. 이 문서를 읽는 방법")
    add_para(doc, "이 문서는 ‘현대백화점의 내부 영업 비밀을 확인했다’고 주장하는 자료가 아닙니다. 공개적으로 확인할 수 있는 운영 사실과, 현재 프로젝트 코드가 이미 가정하고 있는 규칙, 그리고 실운영에서 검증해야 할 설계 제안을 한 문서 안에서 구분해 놓은 실무형 작업 초안입니다.")
    add_table(doc, ["표기", "의미", "사용 규칙"], [
        ("[공개 사실]", "현대백화점·정부·공식 페이지에서 확인된 내용", "실제 운영의 외부 근거로 사용하되 SKU별 내부 정책으로 확대 해석하지 않음"),
        ("[프로젝트 사실]", "현재 저장소의 코드·문서·화면에서 확인된 구조", "목업이므로 실데이터·실행 연동으로 오인하지 않음"),
        ("[제안]", "이 프로젝트를 실운영에 연결할 때의 설계·운영 권고", "담당자·재무·물류·법무가 승인할 정책 후보"),
        ("[추정/가정]", "공개자료가 없거나 시뮬레이션을 위해 둔 값", "원천 데이터 수집 전에는 가격·원가·수요·폐기 확률로 확정하지 않음"),
        ("[하드 차단]", "경제성보다 먼저 막아야 하는 조건", "식품안전·소유권·법규·데이터 품질·물류용량 위반 시 추천/승인 후보에서 제외"),
    ], [1500, 3500, 4360], first_col_bold=True)
    add_callout(doc, "가장 중요한 결론", "현재 프로젝트의 ‘AI 비용’은 아직 계산되지 않습니다. 현재 시뮬레이션의 판매·할인·배송·반품·보관·폐기 비용은 재고 처리 전략의 실행 원가이고, AI 비용은 그 전략을 탐지·추천·시뮬레이션·감사하는 데 드는 별도 비용입니다. 실운영에서는 둘을 합치되, 손익 브리지에서는 반드시 분리해야 합니다.", fill="FFF8E7", accent=AMBER)

    heading(doc, 1, "1. 프로젝트 구조와 시스템 이해")
    heading(doc, 2, "1.1 제품의 실제 목적")
    add_para(doc, "프로젝트의 제품 정의는 소비자용 쇼핑몰이 아니라 현대백화점그룹 계열사 재고 담당자가 자사 악성재고를 처리하기 위한 B2B 의사결정 지원 서비스입니다. 핵심 목표는 ‘더 많이 파는 것’이 아니라 계열사 기준선보다 큰 위험조정 증분 기여현금이익을 만드는 것입니다.")
    add_bullet(doc, "1차 Anchor는 현대백화점 더현대 서울의 2F 여성패션, 3F 남성·잡화, B1 식품관, 1F 뷰티·리빙입니다.")
    add_bullet(doc, "직매입 재고는 소유권·손익·처분권을 현대백화점이 직접 가지므로 P0 핵심 대상입니다. 특약매입·임대매장은 가격·처분 권한을 확인하기 전 직접 할인 후보로 취급하지 않습니다.")
    add_bullet(doc, "본사는 일반 전략을 일괄 승인하는 통제자가 아니라 계열사 간 비용분담·공동 프로모션·그룹 위험·예외 예산을 조정하는 거버넌스 사용자입니다.")
    add_internal_ref(doc, "docs/project-brief.md", "계열사 중심 AI 재고 최적화 서비스 개요")
    add_internal_ref(doc, "docs/decision-policy.md", "기준선·증분 기여현금이익·하드 차단 정책")

    heading(doc, 2, "1.2 현재 저장소에서 확인된 기술 상태")
    add_table(doc, ["영역", "현재 구현", "실운영 해석", "실운영 전 보강"], [
        ("앱", "Next.js 15 / React 19 기반 App Router 프론트엔드", "화면·상태·시뮬레이션 UX 검증용", "API·권한·감사로그·원천 데이터 연결"),
        ("데이터", "src/lib/mock-data.ts의 20개 더현대 서울 품목과 전략 케이스", "샘플 입력·데모 수치", "ERP/POS/WMS/OMS/멤버십/정산 원천과 배치·실시간 동기화"),
        ("AI 파이프라인", "전략 생성 화면의 1~4단계 timed progress", "실제 모델·외부 호출 없음", "규칙 엔진·수요 모델·검색·LLM 설명·모델/데이터 버전 기록"),
        ("시뮬레이션", "src/lib/simulation.ts의 결정론적 판매 반응·비용 계산", "계산 UX와 정책 가드레일 검증용", "실제 비용·수요 반응·불확실성·AI 호출 비용 분리"),
        ("승인·실행", "클라이언트 상태와 모의 Teams/Slack 검토요청", "승인 전 실제 변경 없음", "권한 승인, 가격·재고 mutation API, 실행 결과 회수, 실패 복구"),
        ("분석", "카테고리·가중치·폐기 전후 목업 차트", "설명용 대시보드", "실제 기간·표본·대조군·신뢰구간·비용 대비 효과"),
    ], [1500, 2900, 2500, 2460], first_col_bold=True)
    add_internal_ref(doc, "src/lib/simulation.ts", "SimulationControls·SimulationResult·simulateOption·fallback")
    add_internal_ref(doc, "src/app/strategy/generate/page.tsx", "4단계 생성 화면은 타이머 기반 목업")
    add_internal_ref(doc, "src/app/strategy/[id]/simulate/simulate-client.tsx", "비용 브레이크다운·조정·승인 요청 UI")

    heading(doc, 2, "1.3 현재 업무 흐름")
    add_table(doc, ["단계", "현재 화면/코드", "업무 의미", "AI 비용이 발생하는 지점"], [
        ("1. 탐지", "/inventory/all, /inventory/risk, ProductDetailModal", "위험등급·소유권·보관일·잔여일 확인", "데이터 집계·위험 점수 계산·재분석 호출"),
        ("2. 전략 생성", "/strategy/generate", "원가·보관비·유통기한·과거 반응을 수신한다고 설명", "피처 생성·수요 추정·유사상품 검색·LLM/모델 호출"),
        ("3. 대안 비교", "/strategy/[id]", "순마진·완판·매출 전략 카드 비교", "전략별 시나리오 재계산·설명 생성"),
        ("4. 세밀 조정", "/strategy/[id]/simulate", "할인·쿠폰·포인트·배송·기간·수량·번들·광고·반품·운영비 조정", "조정마다 재계산; 실제 AI 재호출 여부를 제어해야 함"),
        ("5. 담당자 검토", "Teams/Slack 모의 검토요청", "계열사 담당자/협업자 검토와 마감일 설정", "human-in-the-loop 시간·메시지·감사로그 비용"),
        ("6. 승인·실행", "MOCK_OPTIMIZATION_CASES status", "APPROVED→EXECUTING→FINISHED", "승인 후 실행 연계·실패/재시도·결과 회수"),
        ("7. 학습", "/strategy/execution, /analytics", "AI 예상 vs 실제 판매·소진일·기여이익 비교", "평가·리포트 생성·모델 재학습·드리프트 감시"),
    ], [900, 2250, 3300, 2910], first_col_bold=True)
    add_callout(doc, "비용 관리 포인트", "시뮬레이션 슬라이더를 움직일 때마다 LLM을 다시 호출하면 비용과 지연이 폭증합니다. 원칙은 ‘규칙·수식으로 즉시 재계산, AI는 새로운 정보·새 전략·예외 설명이 필요할 때만 호출’입니다.", fill=GREEN_LIGHT, accent=GREEN)

    heading(doc, 2, "1.4 프로젝트 Mock 재고의 기준선 숫자")
    add_para(doc, "다음 수치는 현재 src/lib/mock-data.ts의 20개 품목을 읽어 계산한 데모 기준선입니다. 현대백화점의 실제 재고·매출·폐기 실적이 아니며, 실운영 데이터를 연결할 때 동일한 지표 정의로 교체해야 합니다.")
    add_table(doc, ["지표", "현재 Mock 값", "해석"], [
        ("품목 수 / 수량", "20개 / 3,380개", "4개 층의 단일 점포 데모 범위"),
        ("정상 판매가 총액", "₩588,655,000", "수량×판매가의 표시 가치; 순매출 아님"),
        ("취득원가 총액", "₩260,165,000", "수량×원가; 손익·현금 기준과 분리 필요"),
        ("위험 품목", "14개 / 2,125개 / 판매가 기준 ₩450,095,000", "WARNING·CRITICAL_NEAR·DEAD_STOCK 합계; 표시가 기준 76.5%"),
        ("직매입", "15개 / 2,175개 / 판매가 기준 ₩466,095,000", "표시가 기준 79.2%; P0 우선 대상"),
        ("D-7 이하", "3개 / 310개 / 판매가 기준 ₩54,960,000", "대게·애플망고·한우; 즉시 처리 우선"),
        ("예상 폐기비 합계", "₩59,520,000", "품목별 estimatedDisposalCost×수량; 실제 폐기비와 회계손상 분리 필요"),
        ("보관비/일 합계", "₩2,015,100", "holdingCostPerDay×수량; 공간 대체가치·전력·자본비용 미포함"),
    ], [2450, 2400, 4510], first_col_bold=True)
    add_internal_ref(doc, "src/lib/mock-data.ts", "20개 InventoryItem 필드·값을 직접 집계한 데모 기준선")

    heading(doc, 2, "1.5 현재 시뮬레이션의 비용 구조와 누락")
    add_table(doc, ["현재 필드/계산", "현재 포함", "빠진 비용 또는 위험", "권고"], [
        ("할인·쿠폰·포인트", "grossRevenue에서 차감", "프로모션 중복, 벤더 보전, 포인트 실제 사용률", "명목률과 실제 비용률 분리"),
        ("배송·플랫폼", "온라인 비중×배송지원, 3.2% platformFee", "매장 픽업, 당일배송 capacity, 정산·결제·채널별 수수료", "채널별 실정산 원장 연결"),
        ("피킹·포장·번들", "unit packing/assembly", "점포 인력·라벨·재포장·품질검수·폐기 포장", "활동기준원가(ABC)로 계량"),
        ("반품", "expectedRevenue×returnRate", "반품품 등급하락·재포장·회수·환불시점·브랜드 영향", "반품 상태 전이 모델"),
        ("보관·폐기 회피", "campaignDays·storage, avoidedDisposalCost", "공간 대체가치·냉장전력·폐기 증빙·기부/회수 잔존가치", "회피비용과 잔존가치 명시"),
        ("AI 자체 비용", "없음", "추출·피처·검색·모델·토큰·모니터링·사람 검토·실패 손실", "AI 비용 레이어 신설"),
        ("불확실성", "confidenceScore 단일 점수", "하방확률·신뢰구간·표본·편향·드리프트", "보수/기본/낙관 시나리오"),
    ], [2050, 1850, 3250, 2210], first_col_bold=True)

    heading(doc, 1, "2. 현대백화점 운영 사실 조사")
    heading(doc, 2, "2.1 공개적으로 확인되는 회사·점포·보고 범위")
    add_para(doc, "현대백화점의 2024 지속가능경영보고서 공시는 전국 현대백화점 15개 점포와 아울렛 9개 점포를 보고 범위로 제시합니다. 2024년 1월~12월 성과를 담고 일부 2025년 상반기 성과를 포함하며, 재무정보는 K-IFRS 연결 기준, 비재무 정보는 현대백화점·한무쇼핑 전 사업장 기준이라고 명시합니다. 이는 프로젝트의 ‘더현대 서울 단일관제’가 그룹 전체 운영을 대표하지 않는다는 중요한 경계입니다.")
    add_source_ref(doc, "HDS-ESG-2024")
    add_table(doc, ["공개 확인 항목", "확인된 내용", "프로젝트 적용"], [
        ("보고 경계", "백화점 15개·아울렛 9개, 최근 3개년 추세 데이터", "점포·점포·카테고리별 데이터 권한과 집계 경계를 분리"),
        ("지속가능성 기준", "GRI 2021 핵심 부합, SASB 산업표준 고려, K-IFRS 재무 일치", "재고·폐기·에너지 KPI를 재무·ESG 두 장부로 관리"),
        ("환경경영", "2024 ISO14001 인증, 2025 KCGS 통합 A+·CDP A 공개", "폐기·포장 전략이 브랜드·ESG 리스크에도 영향을 줌"),
        ("사업 채널", "오프라인 점포와 더현대Hi·현대식품관 투홈 등 디지털 접점", "재고 처리 채널별 가격·배송·정산·고객 기대를 분리"),
    ], [1900, 4050, 3410], first_col_bold=True)
    add_source_ref(doc, "HDS-ESG-PAGE")

    heading(doc, 2, "2.2 더현대 서울의 브랜드·방문·팝업 특성")
    add_para(doc, "현대백화점그룹은 더현대 서울 5주년 자료에서 2025년 매출 1조2,864억원, 2030 고객 매출 비중 58%, 팝업스토어 진행 약 1,892건, 사운즈 포레스트 평균 체류 37분을 공개했습니다. 수치는 공식 홍보자료의 집계 범위를 그대로 사용해야 하며, 개별 팝업의 매출·재고회전·수익성을 의미하지는 않습니다.")
    add_source_ref(doc, "HDS-FIVEYEAR")
    add_bullet(doc, "시사점: 단순한 할인보다 ‘발견·경험·한정성’을 이용한 제한 노출, 큐레이션, 팝업 전환이 브랜드 보호형 재고 처리에 적합할 가능성이 큽니다. 이는 공개 성과에서 도출한 제안이며 SKU별 효과는 실험으로 검증해야 합니다.")
    add_bullet(doc, "팝업은 짧은 기간과 한정 물량 때문에 수요를 모으기 쉽지만, 전용 공간·스태프·설치·철거·재고 이동비를 별도 원가로 넣어야 합니다.")

    heading(doc, 2, "2.3 현대퓨처넷과 AI/디지털 접점")
    add_para(doc, "현대백화점은 현대퓨처넷과 협업해 사내 생성형 AI 멘토 ‘Hai’를 도입했고, 13개 직무 139명의 전문가 검증 경험을 학습시켰으며 일평균 350건 이상 질문에 답한다고 공개했습니다. 고객용 쇼핑 AI ‘헤이디’는 브랜드·식당·팝업·전시·프로모션 등 점포 실시간 정보를 연동해 대화형 추천을 제공하며, 현대퓨처넷이 실시간 점포 정보를 생성형 AI가 활용하도록 시스템을 구축했다고 설명합니다.")
    add_source_ref(doc, "HDS-HAI")
    add_source_ref(doc, "HDS-HEYDI")
    add_para(doc, "프로젝트 적용 제안: 현대퓨처넷은 단순 LLM 호출 공급자보다 ERP·인프라·정보보안·메시징·사이니지·전시·운영데이터 분석을 묶는 ‘운영 연결 계층’으로 정의하는 것이 현실적입니다. 재고 전략 AI는 내부 담당자용 Hai 유형의 지식·설명 계층과, 헤이디 유형의 실시간 점포 정보 계층을 분리해 설계해야 합니다.")
    add_source_ref(doc, "FUTURENET-JOB")

    heading(doc, 2, "2.4 실제로 공개된 판매·배송·고객 처리 접점")
    add_table(doc, ["접점", "공식 페이지에서 확인되는 운영", "재고 처리 전략에 넣을 변수"], [
        ("식품 구독", "근거리 배송, 배송일 2일 전 변경, 매장 결제 또는 앱 비대면 결제, 중간 취소는 매장 방문·전체 취소·재결제", "배송 가능 지점, 구독 슬롯, 주문 lead time, 취소·재고 예약, 냉장비"),
        ("선물하기", "구매 후 3일 이내 배송지 입력, 7일 이내 구매점 방문수령, 미수령 시 받지 않기·구매점 취소", "선물 재고 예약, 미수령 확률, 배송·취소·재판매 시점"),
        ("교환·환불", "일반적으로 구입 후 7일 이내, 미사용·택 부착·영수증, 동일 결제수단; 신선식품·냉장·냉동·베이커리 등 예외", "반품률이 아니라 반품 가능성·재판매 가능성·상품군 예외를 함께 계산"),
        ("식품 안전", "식품안전 보증제와 식품 감식관, 위생·품질 이상 시 교환·환불 및 보상", "식품 임박재고는 수익보다 안전·품질 데이터가 우선 하드 차단"),
    ], [1700, 4400, 3260], first_col_bold=True)
    add_source_ref(doc, "HDS-SUBSCRIPTION")
    add_source_ref(doc, "HDS-SERVICE")
    add_source_ref(doc, "HDS-RETURNS")

    heading(doc, 2, "2.5 VIP·H.Point·친환경 참여의 경제성")
    add_para(doc, "현대백화점 VIP 프로그램은 마일리지·리워드·전용 할인·라운지 등 고객 가치가 높은 대신, 모든 상품군에 동일하게 적용되지 않습니다. 2026 Club YP 안내에는 임대매장, 해외 명품, 가전, 식품, 식당가, 푸드코트, 할인매장 등이 할인 제외로 명시되어 있고, 온라인몰도 혜택별로 제외될 수 있습니다. 따라서 ‘VIP 타깃’은 만능 할인 수단이 아니라 상품군·채널·회원등급별 조건을 확인한 제한적 수요 자극 수단입니다.")
    add_source_ref(doc, "HDS-VIP-2025")
    add_source_ref(doc, "HDS-YP-2026")
    add_source_ref(doc, "HPOINT")
    add_table(doc, ["수단", "장점", "비용·위험", "AI 전략에서의 사용 조건"], [
        ("VIP 큐레이션", "가격 공개 범위를 제한하고 구매력이 높은 고객에 노출", "혜택 제외·고객 피로·정상 판매 잠식", "상품군·등급·채널 eligibility 통과 시에만"),
        ("H.Point 적립", "할인율을 낮춰 표시가를 방어하며 재방문 유도", "실제 사용률·소멸·정산비용이 명목 적립률과 다름", "명목 포인트율×예상 사용률×자금/정산비로 계산"),
        ("친환경 참여 리워드", "브랜드·ESG 메시지와 자원순환을 결합", "재고 판매를 직접 보장하지 않음", "제품 판매가 아닌 포장 회수·캠페인 비용/효과로 분리"),
        ("팝업/체험", "체류·발견·한정성으로 할인 의존 완화", "공간·스태프·설치·철거·기회비용", "예상 매출보다 증분 기여현금이익과 공간 대체가치로 판단"),
    ], [1600, 2700, 2800, 2260], first_col_bold=True)

    heading(doc, 2, "2.6 식품·명절·신선상품의 실제 계획 신호")
    add_para(doc, "현대백화점은 2025년 추석 한우 선물세트를 11만 세트로 전년 대비 10% 확대하고, 산소치환 포장을 적용한다고 공개했습니다. 이는 실제 운영이 ‘재고를 일괄 보유한 뒤 할인’만이 아니라, 수요·시즌·포장·선도 유지·배송을 함께 계획한다는 공개 사례입니다. 프로젝트에서는 시즌성 품목의 입고 수량, 예약 판매, 선물 배송일, 포장 방식, 잔여 처리 가능성까지 입력해야 합니다.")
    add_source_ref(doc, "HDS-BEEF-2025")
    add_bullet(doc, "[제안] 명절·예약 상품은 판매기간 전체의 총수요보다 ‘남은 배송일별 확정 주문+예상 추가주문+취소/미수령’을 분리하는 일자별 재고예약 모델이 필요합니다.")
    add_bullet(doc, "[하드 차단] 소비기한·보관조건·선도검사 데이터가 없으면 AI가 할인율을 추천하기 전에 판매·기부·폐기 가능 여부를 담당자가 확인해야 합니다.")

    heading(doc, 2, "2.7 자원순환·폐기·브랜드 이미지")
    add_para(doc, "현대백화점은 점포에서 발생한 비닐을 분리배출→수집·보관→압축·보관→열분해유 생산→정제→원료화해 새 비닐로 활용하는 프로세스를 공개했고, Project100은 점포 폐지를 재생지로 원료화해 쇼핑백·명절 포장재로 사용합니다. 이는 ‘상품 폐기’의 처리 경로를 직접 공개한 것이 아니라 점포 운영 폐기물·포장재의 자원순환 사례입니다.")
    add_source_ref(doc, "HDS-CIRCULAR")
    add_source_ref(doc, "HDS-PROJECT100")
    add_callout(doc, "중요한 분리", "상품 손실(유통기한 경과·파손·시즌 이월)과 포장 폐기물(폐비닐·박스·쇼핑백)을 한 비용으로 합치지 마십시오. 상품은 판매·반품·기부·벤더회수·폐기 의사결정이고, 포장재는 자원순환·ESG·운영비 의사결정입니다. 둘을 연결하되 원가 원장은 분리해야 합니다.", fill="FFF8E7", accent=AMBER)

    heading(doc, 1, "3. 현대백화점 재고 처리 전략: 공개 사실과 실무형 설계")
    heading(doc, 2, "3.1 공개자료로 알 수 없는 것과 반드시 요청해야 하는 것")
    add_para(doc, "현대백화점의 공개 페이지·보고서에서 SKU별 일일 판매속도, 직매입/특약매입 계약별 할인 승인권, 점포별 폐기·벤더회수·기부 비율, 채널별 물류단가, 실제 마크다운 알고리즘은 확인되지 않습니다. 따라서 아래 전략은 공개된 서비스 접점과 프로젝트 정책을 이용한 실무 설계안이며, 도입 전에 현대백화점 담당자·현대퓨처넷·재무·물류·법무가 원천 데이터를 확인해야 합니다.")
    add_table(doc, ["필수 요청 데이터", "필드 예시", "수집 주기", "미수집 시 차단"], [
        ("소유권·계약", "purchaseType, vendor, settlement, return right, markdown approval", "입고/계약 변경", "특약·임대매장 직접 할인 금지"),
        ("판매·예약", "POS/OMS 판매, 예약, 취소, 미수령, 채널·시간대", "일/시간", "수요 예측 신뢰도 하향"),
        ("품질·기한", "lot, 소비기한, 보관조건, 검사, 리콜·폐기 사유", "입고/일", "식품 전략 하드 차단"),
        ("물류", "점포·센터 재고, pick/pack, 배송/픽업 capacity, 냉장·냉동", "일/실시간", "배송 약속 전략 금지"),
        ("비용·정산", "원가, 수수료, 할인 부담주체, 포인트 사용률, 반품비, 폐기비", "월/거래", "증분이익 산출 불가"),
        ("브랜드·고객", "VIP eligibility, campaign fatigue, brand markdown policy", "캠페인별", "공개 타임세일·가격 훼손 전략 금지"),
    ], [2000, 3550, 1700, 2360], first_col_bold=True)

    heading(doc, 2, "3.2 처리 우선순위: 만료·소유권·가치·브랜드 순")
    add_para(doc, "최대 이윤은 ‘마진율이 가장 높은 상품부터’가 아니라, 처리기한을 넘길 때 발생하는 하방 손실과 AI·운영 비용까지 고려해 결정해야 합니다. 다음 우선순위는 프로젝트의 현재 정책과 현대백화점 공개 운영 특성을 결합한 권고입니다.")
    add_number(doc, "D-7 이하 신선·냉장·냉동 직매입: 안전·품질 승인 후 당일/익일 판매·예약·근거리 배송·회수 가능성을 비교합니다. 남은 시간이 짧으므로 비싼 LLM 설명을 반복하지 말고 규칙 엔진과 담당자 승인으로 빠르게 움직입니다.")
    add_number(doc, "D-8~30 식품·디저트·밀키트: 선도·배송·구독·선물·타임세일·묶음 조합을 일 단위로 재예측합니다. 기부·폐기 cutoff와 판매 가능 cutoff를 별도로 둡니다.")
    add_number(doc, "시즌 D-14~90 패션: 정상가 공개 범위를 보호하면서 VIP/앱 큐레이션·한정 팝업·점포 내 위치 이동을 먼저 검토하고, 목표 미달 시 아울렛·별도 채널·제한 타임세일로 단계 전환합니다.")
    add_number(doc, "비기한 프리미엄 장기재고: 공간 대체가치와 자본비용이 높더라도 즉시 대폭 할인하지 않고, 브랜드 적합 채널·큐레이션·예약·고객 맞춤형 노출을 우선합니다.")
    add_number(doc, "특약매입·임대매장: 계약상 소유권·처분권·할인부담·반품 조건이 확인되기 전에는 AI가 가격을 제안하지 않고 협의 과제로 보냅니다.")

    heading(doc, 2, "3.3 카테고리별 처리 플레이북")
    add_table(doc, ["카테고리", "1차 전략", "2차 fallback", "주요 비용", "하드 차단"], [
        ("신선식품·수산·과일", "당일 타임세일·예약/선물 배송·근거리 배송", "기부/회수/폐기 cutoff 비교", "선도검사·냉장·배송·폐기", "소비기한·보관·검사 불충족"),
        ("밀키트·디저트", "묶음·선물·근거리 구독·앱 타깃", "제한 할인·기부/폐기", "포장·냉장·반품 불가·폐기", "표시·알레르기·소비기한"),
        ("시즌 패션", "VIP/앱 큐레이션·팝업·소량 할인", "아울렛·별도 채널·단계적 타임세일", "할인·공간·인력·정상판매 잠식", "브랜드 가격정책·소유권"),
        ("가방·시계·주얼리", "프리미엄 큐레이션·예약·고객 1:1", "승인된 별도 채널·벤더 협의", "자본·보안·보험·정품/보증", "가품·보증·가격 승인"),
        ("리빙·향·식기", "선물·사용목적 번들·공간 큐레이션", "아울렛·시즌 기획전", "포장·파손·재고 이동·반품", "파손·번들 적합성"),
        ("특약·임대", "벤더/브랜드와 공동 캠페인", "반품·회수·협의 후 제한 판매", "수수료·분담·정산·권한", "계약 미확인"),
    ], [1700, 2500, 2400, 1750, 1010], first_col_bold=True, font_size=8.35)

    heading(doc, 2, "3.4 현재 Mock 위험재고에 대한 권고 매핑")
    add_para(doc, "아래는 현재 프로젝트 Mock 재고에 대해 위 플레이북을 적용한 우선순위 예시입니다. 가격·판매량은 실운영 추천값이 아니며, 어떤 데이터를 추가해야 전략이 성립하는지 보여주는 용도입니다.")
    add_table(doc, ["품목", "상태", "즉시 권고", "fallback", "AI/운영 비용 측정"], [
        ("INV-THS-015 대게 D-3", "DEAD_STOCK / 직매입", "선도검사 후 당일 한정 판매·예약배송", "기부/폐기 cutoff", "당일 rule run, 냉장·포장·폐기비, 담당자 분 단위"),
        ("INV-THS-013 애플망고 D-5", "DEAD_STOCK / 직매입", "숙도·품질 확인 후 소량 타임세일·세트", "폐기/회수 비교", "숙도 판정·폐기확률·가격 반응"),
        ("INV-THS-012 한우 D-6", "DEAD_STOCK / 직매입", "식품팀 승인 후 퇴근 시간대·예약·배송", "마감 전 단계 할인", "냉장·배송 capacity·미수령·반품"),
        ("INV-THS-014 밀키트 D-9", "CRITICAL_NEAR / 직매입", "구독·다수 구매·식품관 타깃 번들", "제한 할인·기부", "SKU·lot·알레르기·포장 비용"),
        ("INV-THS-001 캐시미어 D-14", "DEAD_STOCK / 직매입", "2F 제한 큐레이션·팝업·타깃 15% 수준 테스트", "앱 핫딜→아울렛", "브랜드 잠식·VIP eligibility·공간 대체가치"),
        ("INV-THS-010 초콜릿 D-18", "CRITICAL_NEAR / 직매입", "선물·디퓨저 등 목적 적합 번들", "앱 한정·식품관 타임세일", "번들 조립·선물포장·소비기한"),
        ("INV-THS-006 구스 베스트 D-20", "CRITICAL_NEAR / 직매입", "3F 시즌 큐레이션·앱 타깃", "아울렛·단계 할인", "시즌 수요·사이즈 분포·이월 감가"),
        ("INV-THS-002 원피스 D-30", "CRITICAL_NEAR / 직매입", "여성패션 팝업·H.Point 타깃", "30% 타임세일 전환", "할인 탄력·사이즈별 잔여"),
        ("INV-THS-003 가죽 자켓 D-45", "WARNING / 직매입", "프리미엄 1:1·가을 연계 큐레이션", "승인된 아울렛/별도 채널", "브랜드 가격·보안·반품"),
        ("INV-THS-018 디퓨저 D-60", "WARNING / 직매입", "초콜릿·와인·선물 목적 번들", "리빙 기획전", "번들 적합성·파손·포장"),
        ("INV-THS-005 브리프케이스 D-90", "WARNING / 직매입", "3F 라운지·앱 타깃, 기존 CASE-2026-003 학습", "노출 보강 후 아울렛", "예측오차·광고·전환율"),
        ("INV-THS-017 오크 디퍼 D-180", "DEAD_STOCK / 직매입", "프리미엄 리빙 큐레이션·공간 대체가치 비교", "아울렛·선물 번들", "부피·파손·장기 보관"),
        ("INV-THS-020 아이세트 D-200", "SAFE / 임대", "소유자·할인 권한 확인 전 유지", "벤더 공동행사", "정산·분담·반품 조건"),
    ], [2500, 1450, 2350, 1750, 1310], first_col_bold=True, font_size=7.8)
    add_internal_ref(doc, "src/lib/mock-data.ts", "위험 상태·소유권·기한·보관비·폐기비 및 기존 전략 CASE-2026-001~004")

    
    heading(doc, 2, "3.5 현대백화점 패션 부문 대상 Oracle Retail 기반 재고 위험 분석 및 시즌 중 계획 체계")
    add_para(doc, "Oracle Retail의 In-season Planning(시즌 중 상품계획) 표준 체계를 참고하여, 현대백화점 패션 부문 입점 브랜드를 대상으로 한 재고 위험 분석 및 시즌 중 처리 추천 모델을 정의합니다. 전체 상품군에 일관적 고정 기준을 무리하게 적용하는 대신, 패션 부문에 집중하여 카테고리·브랜드·상품태그별 계층적 정책 상속 구조를 설계합니다.")

    add_callout(doc, "패션 부문 한정 이유", "패션 상품은 프로모션·시즌오프·가격인하·점포간 이동·아울렛 이관 전략과 가장 직접적으로 연결됩니다. 식품(소비기한/신선도), 가전(신모델/기술노후화), 화장품(사용기한/단종) 등은 위험 판단 기준과 물류 제약이 전혀 다르므로, 1차 범위를 현대백화점 패션 입점 브랜드로 명확히 제한하여 실무적 타당성을 확보합니다.", fill=GREEN_LIGHT, accent=GREEN)

    add_table(doc, ["Oracle 지표/개념", "현대백화점 패션 실무 정의", "수식 / 시스템 로직", "재고 처리 활용"], [
        ("ROS (Rate of Sales)", "일/주 단위 판매속도", "ROS = Q_sold ÷ t_days", "동일 재고량이라도 ROS가 급감하면 과잉재고 경고"),
        ("Sell Thru % (판매율)", "입고 대비 순판매 비율", "ST% = (Q_sales - Q_returns) ÷ (Q_bOH + Q_receipts)", "목표 판매율 미달 시 행사·할인·아울렛 검토"),
        ("Needed Sales Lift %", "목표 달성을 위한 판매 증분 필요율", "Lift% = (ST%_target - ST%_projected) ÷ ST%_projected", "목표 달성에 필요한 인하/프로모션 강도 결정"),
        ("WOS (Weeks of Supply)", "현재 재고의 판매 소진 예상 주수", "WOS_proj = Q_OH ÷ ROS_weekly", "Target WOS 초과 시 Overage(과잉재고) 경고 발생"),
        ("STD vs BTA", "시즌 누적실적(STD)과 잔여목표(BTA)", "STD = Season To Date, BTA = Balance To Achieve", "Trend Review 3단계(Keep / Add / Drop) 판단"),
        ("계획 버전 (WP/CP/OP)", "수정중(WP), 승인(CP), 최초(OP)", "WP(Working Plan) → MD승인 → CP(Current Plan)", "AI 시뮬레이션 결과를 WP에 반영 후 최종 승인"),
    ], [2200, 2400, 2400, 2360], first_col_bold=True, font_size=8.4)

    heading(doc, 3, "계층적 정책 상속 구조 (Hierarchical Policy System)")
    add_para(doc, "현대백화점의 수많은 입점 브랜드와 SKU의 목표 판매율 및 WOS를 개별 수동 입력하지 않고, 상위 수준 기본값부터 상품별 예외까지 자동 상속·보정하는 5단계 구조를 적용합니다.")
    add_table(doc, ["단계", "계층", "예시 설정값", "우선순위"], [
        ("1단계", "패션 상품군 공통 기본값", "목표 Sell Thru 65%, 목표 WOS 5주, 장기재고 기준 90일", "최하위 (Default)"),
        ("2단계", "카테고리별 기준", "아우터(ST 80%, WOS 3주), 기본셔츠(ST 60%, WOS 6주)", "카테고리 보정"),
        ("3단계", "브랜드별 정책 (brand_policy)", "브랜드A(아울렛 이관 가능), 브랜드B(할인율 10% 제한)", "브랜드 계약 준수"),
        ("4단계", "상품 태그/속성 보정", "SEASONAL/TREND(+5%p ST), BASIC(-5%p ST), LIMITED", "속성별 위험 보정"),
        ("5단계", "상품/점포 직접 오버라이드", "한정판 패딩 (ST 90%, 할인 불가, 아울렛 불가)", "최우선 (Override)"),
    ], [1200, 2600, 4200, 1360], first_col_bold=True, font_size=8.4)

    add_source_ref(doc, "ORACLE-RETAIL-INSEASON")

    heading(doc, 1, "4. AI 비용관리 전략: 재고 처리 AI의 진짜 원가")
    heading(doc, 2, "4.1 AI 비용과 재고 처리 비용을 분리하는 이유")
    add_para(doc, "현재 프로젝트의 할인·쿠폰·배송·포장·반품·보관·폐기 회피는 ‘선택한 재고 처리 전략’의 비용입니다. 반면 AI 비용은 위험재고를 발견하고, 판매·폐기 확률을 추정하고, 전략 후보를 생성·설명하고, 시뮬레이션·감사·학습하는 데 드는 비용입니다. 재고 처리 결과에서 AI 비용을 누락하면 모든 AI 추천이 실제보다 이익이 큰 것처럼 보입니다.")
    add_table(doc, ["원가 층", "예시", "현재 프로젝트 대응"], [
        ("재고 매몰원가", "이미 취득한 원가·장부가", "기준선/증분이익에서 전략으로 변하지 않는 부분은 분리"),
        ("전략 실행원가", "할인·쿠폰·포인트·배송·피킹·포장·반품·보관·폐기", "현재 SimulationResult에 대부분 존재"),
        ("AI 결정원가", "데이터 추출·특징 생성·검색·모델·LLM·도구·오케스트레이션", "신설 필요"),
        ("AI 운영원가", "모니터링·평가·재학습·보안·로그·인프라·지원", "신설 필요"),
        ("AI 실패원가", "잘못된 할인·소진 실패·브랜드 훼손·규제·기회비용", "하방 시나리오·reserve로 반영"),
        ("사람 비용", "MD 검토·승인·물류/재무/법무 협의", "human-in-the-loop 시간으로 계량"),
    ], [1800, 3750, 4060], first_col_bold=True)
    add_callout(doc, "손익 공식", "AI로 조정된 순가치 = (AI 전략의 위험조정 기여현금이익 − 계열사 기준선 기여현금이익) − AI 결정원가 − 추가 사람·운영원가 − AI 실패 reserve. 음수이면 판매 매출이 커도 ‘AI가 만든 경제적 가치’는 음수일 수 있습니다.", fill="FFF8E7", accent=AMBER)

    heading(doc, 2, "4.2 권장 AI 비용 원장")
    add_table(doc, ["원가 코드", "정의", "계산식/측정", "귀속 기준"], [
        ("DATA", "원천 데이터 추출·정제·조인", "작업시간·배치/ETL 비용·행/GB", "계열사·점포·실행일"),
        ("FEATURE", "판매속도·기한·가격·유사상품 피처", "CPU/작업시간·피처 재생성 횟수", "케이스·품목"),
        ("SEARCH", "유사상품·정책·계약·FAQ 검색", "쿼리 수·embedding·vector storage", "케이스·검색 인덱스"),
        ("MODEL", "수요·할인탄력·폐기확률 모델", "학습·추론 호출 수·GPU/CPU 시간", "모델 버전·케이스"),
        ("LLM", "설명·전략 서술·예외 요약", "입력/출력/캐시 토큰×가격", "모델·프롬프트 버전·케이스"),
        ("TOOL", "ERP/POS/WMS/메시징/API 호출", "호출 수·실패/재시도·외부 API 단가", "도구·케이스"),
        ("ORCH", "워크플로우·큐·서버·캐시", "실행시간·메모리·스토리지·로그 GB", "실행 ID"),
        ("HUMAN", "MD·물류·재무·법무 검토", "검토분×표준 시간당 원가", "승인자·예외 종류"),
        ("EVAL", "정확도·편향·드리프트·대조군 평가", "평가 케이스 수·라벨링 시간", "모델/릴리스"),
        ("RESERVE", "AI 오류로 인한 하방 손실 대비", "오류확률×오류 영향×귀속률", "위험등급·카테고리"),
    ], [1200, 2700, 3300, 2410], first_col_bold=True, font_size=8.4)

    # Keep the formula callout and its metric table together; the preceding
    # ledger is dense enough to fill the page on Korean rendering.
    doc.add_page_break()
    heading(doc, 2, "4.3 케이스별 AI 원가 공식")
    add_para(doc, "권장 계산식은 아래와 같습니다. 숫자 단위는 원화이며, 실제 공급자·인프라·인건비 계약에 맞게 파라미터화합니다.")
    add_callout(doc, "케이스 AI 원가", "C_AI_case = C_DATA + C_FEATURE + C_SEARCH + C_MODEL + C_LLM + C_TOOL + C_ORCH + C_HUMAN + C_EVAL + C_SHARED. C_SHARED는 월 고정 인프라·보안·운영비를 케이스 수 또는 사용량 기준으로 일관되게 배부합니다.", fill=GREEN_LIGHT, accent=GREEN)
    add_para(doc, "AI 비용을 품목 수로 나눌 때는 단순히 ‘전체 재고 수’가 아니라 실제 전략 의사결정 대상 품목과 시뮬레이션 횟수를 사용합니다. 예를 들어 동일 케이스의 할인율 슬라이더 100회 조정은 LLM 100회가 아니라 수식 재계산 100회로 처리하고, 새로운 데이터나 전략 후보가 생성될 때만 모델/LLM 비용을 추가합니다.")
    add_table(doc, ["지표", "식", "경영 판단"], [
        ("AI 비용/케이스", "C_AI_case", "케이스 승인 전에 고정·변동 원가 확인"),
        ("AI 비용/품목", "C_AI_case ÷ 대상 SKU 수", "대량 일괄 처리와 단건 예외 비교"),
        ("AI 비용/처리수량", "C_AI_case ÷ 예상 소진 수량", "신선식품 긴급 처리와 장기재고 비교"),
        ("AI 비용/₩1백만원 가치", "C_AI_case ÷ max(순 AI 가치/1,000,000, 1)", "AI가 비싼 케이스를 차단"),
        ("AI ROI", "(순 AI 가치 − C_AI_case) ÷ C_AI_case", "계열사·점포·모델별 투자 효율"),
        ("손익분기 가치", "C_AI_case ÷ 귀속률", "AI가 만들어야 하는 최소 증분 가치"),
    ], [2200, 4300, 3110], first_col_bold=True)

    heading(doc, 2, "4.4 현재 코드 필드에 추가할 AI 비용 입력")
    add_table(doc, ["추가 필드", "의미", "기본 입력/범위", "어느 화면에 보일지"], [
        ("aiRunCount", "한 케이스에서 새 AI 추론이 발생한 횟수", "실행 로그에서 집계", "전략 상세·실행 성과"),
        ("aiInputTokens / aiOutputTokens", "LLM 입력·출력 토큰", "공급자 usage 응답", "고급 비용 패널"),
        ("cachedInputTokens", "캐시된 입력 토큰", "usage 응답", "비용 절감률"),
        ("retrievalQueryCount", "정책·계약·유사상품 검색 횟수", "쿼리 로그", "비용 브레이크다운"),
        ("dataPrepMinutes", "데이터 정제·검증 시간", "작업 로그", "인건비 원장"),
        ("humanReviewMinutes", "MD·물류·재무·법무 검토 시간", "승인 이벤트", "실제 기여이익"),
        ("aiInfraAllocatedCost", "공유 인프라 케이스 배부", "월비용÷사용량", "AI 원가"),
        ("aiFailureProbability", "잘못된 추천/실패확률", "검증 데이터·보수적 구간", "하방 reserve"),
        ("aiFailureImpact", "실패 시 추가 손실", "브랜드·폐기·가격·기회비용", "위험조정 가치"),
        ("netAfterAICost", "AI 비용 차감 후 순가치", "시뮬레이션 결과", "최우선 KPI"),
    ], [2300, 3300, 2000, 2010], first_col_bold=True, font_size=8.4)
    add_internal_ref(doc, "src/lib/simulation.ts", "기존 SimulationControls·SimulationResult에 추가할 후보 필드")

    heading(doc, 2, "4.5 모델·호출 라우팅으로 비용을 낮추는 방법")
    add_table(doc, ["업무", "권장 방식", "비용 제어", "품질 가드레일"], [
        ("위험 태깅", "규칙+통계/소형 모델", "매일 배치·이벤트 변경 때만 재실행", "소유권·기한·품질 하드 차단"),
        ("판매/기한 예측", "시계열·회귀·분류 모델", "CPU 추론·피처 캐시", "표본 부족·신뢰구간 표시"),
        ("유사상품 검색", "embedding+벡터 검색", "상품·정책 임베딩 재사용", "동일 상품 65% 등 프로젝트 가중치는 검증 전 가정"),
        ("전략 후보 생성", "템플릿·수식 우선, LLM은 설명/예외", "후보 수 상한·프롬프트 캐시", "금지 조건 먼저 필터"),
        ("슬라이더 시뮬레이션", "클라이언트/서버 수식 재계산", "LLM 재호출 금지", "원가·현재고·기한 읽기 전용"),
        ("MD 메시지", "구조화된 템플릿+소형 모델", "짧은 출력·배치", "사람 승인 전 발송·실행 금지"),
        ("모델 재평가", "비동기 Batch", "야간·주간 batch/캐시", "대조군·예측오차·드리프트"),
    ], [1900, 2750, 2700, 2260], first_col_bold=True, font_size=8.4)
    add_source_ref(doc, "OPENAI-PRICE")
    add_para(doc, "공급자 가격은 예시일 뿐이며 조달 시 최신 가격을 다시 확인해야 합니다. 핵심은 특정 모델 가격을 하드코딩하는 것이 아니라, 모델별 입력·출력·캐시·batch·tool 비용을 모두 usage 원장으로 남기는 것입니다.", italic=True, color=MUTED, size=9.5)

    heading(doc, 2, "4.6 현재 Mock 케이스에 적용한 비용 가드레일 예시")
    add_para(doc, "다음은 ‘AI 자체 원가가 예상 AI 가치의 일정 비율을 넘지 않게 한다’는 운영 규칙의 예시입니다. 실제 예산이 아니며, 데이터가 없는 단계에서 사용할 수 있는 임시 가드레일입니다.")
    add_table(doc, ["케이스", "기존 예상 증분 기여이익", "임시 AI 원가 상한(1%)", "운영 해석"], [
        ("CASE-2026-001 캐시미어", "₩29,120,000", "₩291,200", "장기 패션은 반복 설명보다 1회 분석+수식 조정 우선"),
        ("CASE-2026-004 한우", "₩18,200,000", "₩182,000", "긴급 품질·물류 판단이 핵심; 호출보다 데이터 검증/담당자 속도 우선"),
        ("CASE-2026-002 초콜릿+디퓨저", "₩8,900,000", "₩89,000", "번들 검색·포장·브랜드 적합성 검토를 포함하되 후보 수 제한"),
    ], [2600, 2300, 2200, 2510], first_col_bold=True)
    add_callout(doc, "더 정확한 규칙", "1%는 보편 법칙이 아닙니다. 신선식품·대형 손실 케이스는 AI 비용 비율이 조금 높아도 순손실을 크게 줄일 수 있고, 반대로 단순·반복·저가 케이스는 1%보다 훨씬 낮아야 합니다. 점포·카테고리별 월 예산, 위험등급별 상한, 손익분기 가치로 재설정하십시오.", fill="FFF8E7", accent=AMBER)

    
    heading(doc, 2, "4.7 AI 판단 기준 최적화 수학 모델 및 제약조건 수식")
    add_para(doc, "AI가 단순 할인율이나 매출 최대화를 추천하는 것이 아니라, 현대백화점의 하드 차단 조건(Hard-Stops)을 완벽히 준수하면서 'AI 비용 차감 후 순가치(NetValue)'를 극대화하는 수학적 최적화 문제(Optimization Formulation)를 정의합니다.")

    add_callout(doc, "최적화 목적함수 (Objective Function)", "max_{s ∈ S} NetValue(s) = ΔProfit(s) - C_AI_case(s) - C_human(s) - Reserve_downside(s) | where ΔProfit(s) = Cash_AI(s) - Cash_baseline", fill="FFF8E7", accent=AMBER)

    heading(doc, 3, "하드 차단 조건 (Hard-Stop Constraints)")
    add_para(doc, "다음 5대 제약조건 중 하나라도 위반하는 전략 s는 목적함수 계산 전 후보 집합 S에서 즉시 제거(Pruning)됩니다.")
    add_table(doc, ["제약조건 종류", "수학적 표현 / 조건", "위반 시 조치", "현대백화점 현장 의미"], [
        ("1. 식품안전/유통기한", "t_remaining >= t_safety_cutoff", "하드 차단 (BLOCK)", "소비기한 임박 식품 임의 연장/재포장 금지"),
        ("2. 소유권/아울렛 권한", "OutletAllowed(item) == 1 (Strategy includes Outlet)", "하드 차단 (BLOCK)", "특약매입/브랜드 계약상 아울렛 이관 금지 조항 준수"),
        ("3. 물류/점포 용량", "Q_transfer <= Capacity_target_store", "하드 차단 (BLOCK)", "목적지 점포/아울렛 창고 수용 용량 초과 이동 금지"),
        ("4. 브랜드 할인율 상한", "d_s <= d_max_brand", "하드 차단 (BLOCK)", "브랜드 가치 보호를 위한 최고 할인율 제한"),
        ("5. AI 최소 ROI", "ROI_AI(s) >= ROI_min (예: 100%)", "후보 제외 (EXCLUDE)", "AI 판단 비용보다 증분이익이 작은 비효율 추천 차단"),
    ], [1900, 2900, 2000, 2560], first_col_bold=True, font_size=8.4)

    heading(doc, 1, "5. 실무형 최적 재고 처리 전략")
    heading(doc, 2, "5.1 기준선(Do Nothing)부터 만든다")
    add_para(doc, "기준선은 단순히 ‘아무것도 하지 않음’이 아니라, 프로모션을 하지 않고도 실행 가능한 가장 좋은 운영 대안입니다. 품목마다 정상 판매 속도, 시즌/소비기한 감소, 보관비, 공간 대체가치, 향후 할인·반품·폐기 가능성을 날짜별로 계산합니다.")
    add_table(doc, ["기준선 구성", "필수 값", "주의"], [
        ("수요", "최근 7/30/90일 판매·요일·날씨·시즌·행사", "판매 0은 수요 0이 아니라 노출 부족일 수 있음"),
        ("기한", "소비기한/시즌 종료/품질등급 cutoff", "처리기한은 비용이 아니라 제약조건"),
        ("보관", "일 보관비·냉장전력·공간 대체가치", "고정 임차료는 증분비용과 분리"),
        ("잔존", "기한 이후 할인·아울렛·회수·기부·폐기 확률", "상태 전이는 중복 계산하지 않음"),
        ("정상 판매 잠식", "프로모션으로 정상 판매가 줄어드는 효과", "총매출이 아닌 증분 가치에 반영"),
    ], [2200, 3800, 3610], first_col_bold=True)

    heading(doc, 2, "5.2 시나리오 설계")
    add_table(doc, ["시나리오", "수요 가정", "운영 가정", "선택 기준"], [
        ("보수", "할인 탄력 낮음·반품/미수령 높음", "배송·공간·인력 용량 제한", "하방 손실이 허용범위 내인지"),
        ("기본", "최근 유사상품·동일상품 가중치", "승인된 채널·정상 capacity", "위험조정 증분 기여현금이익 최대"),
        ("낙관", "행사·팝업·VIP 노출이 계획대로 작동", "추가 노출·배송·인력이 확보", "낙관값 단독 승인 금지"),
    ], [1800, 3100, 2600, 2110], first_col_bold=True)
    add_para(doc, "현재 ProductDetailModal과 analytics 화면의 동일상품 65%·유사 카테고리 25%·시즌 10% 가중치는 프로젝트 가정입니다. 운영 전에는 동일상품의 표본 수·시즌 편향·가격 변경 이력·프로모션 노출 차이를 검증하고, 표본이 부족하면 가중치를 자동으로 낮추거나 신뢰도 경고를 띄워야 합니다.")
    add_internal_ref(doc, "src/components/inventory/product-detail-modal.tsx", "위험 분석·기준선·가중치 목업")
    add_internal_ref(doc, "src/app/analytics/page.tsx", "가중치·탄력도·폐기 전후 목업 지표")

    heading(doc, 2, "5.3 Fallback 트리")
    add_table(doc, ["누적 판매 달성률", "권장 조치", "할인·비용 원칙", "승인"], [
        ("90–100%", "현재 전략 유지·일별 검증", "추가 비용을 만들지 않음", "계열사 담당자 모니터링"),
        ("70–89%", "앱/매장 노출 보강·제한 쿠폰", "추가 혜택의 증분 가치가 비용보다 큰지", "계열사 담당자"),
        ("50–69%", "H.Point/타깃 혜택·번들·시간대 조정", "할인·배송·조립·반품 비용 재계산", "계열사 담당자+재무/물류"),
        ("30–49%", "아울렛·별도 채널·강화 타임세일", "브랜드·정상판매 잠식·잔여 폐기 위험 포함", "예외 기준이면 본사 공유"),
        ("0–29%", "판매보다 회피 손실이 작은 회수·기부·폐기", "추가 할인은 자동 실행하지 않음", "담당자 승인; 법규/그룹 위험은 본사"),
    ], [1800, 3300, 2700, 1810], first_col_bold=True)
    add_internal_ref(doc, "src/lib/simulation.ts", "buildSimulationFallback 5구간 규칙")

    heading(doc, 1, "6. 마케팅·물류·폐기·브랜드를 함께 최적화하는 방법")
    heading(doc, 2, "6.1 마케팅")
    add_bullet(doc, "할인 공개 범위를 단계화합니다: 내부/MD 타깃 → VIP·앱 개인화 → 점포 큐레이션/팝업 → 아울렛/별도 채널 → 공개 타임세일.")
    add_bullet(doc, "VIP 혜택 제외 상품군이 있으므로 식품·해외명품·임대매장 등을 VIP 할인으로 가정하지 않습니다. H.Point·선물·예약·콘텐츠 등 다른 혜택으로 재설계합니다.")
    add_bullet(doc, "팝업은 ‘재고를 숨기는 곳’이 아니라 체류·발견·브랜드 스토리와 판매를 연결하는 공간입니다. 행사장 임차·인력·설치·철거·이동비와 정상 공간의 대체가치를 함께 산정합니다.")
    add_bullet(doc, "캠페인 성과는 매출만 보지 않고 신규/재방문, 정상가 판매 잠식, 브랜드 검색량·문의, 재고 소진, 폐기 회피, AI 비용까지 함께 봅니다.")

    heading(doc, 2, "6.2 물류·배송")
    add_bullet(doc, "식품 구독·선물 배송은 지점·배송일·냉장 capacity가 제약입니다. ‘배송 가능 상품’과 ‘오늘 배송 가능한 수량’을 분리해 재고 예약을 관리합니다.")
    add_bullet(doc, "온라인 비중이 올라갈수록 포장·피킹·배송·반품·미수령 비용이 커질 수 있으므로 무료배송은 예상 증분가치의 일정 비율을 넘지 않게 제한합니다.")
    add_bullet(doc, "점포 간 이동은 P2로 두되, 이동 전후 가격·소유권·정산·재고 정확도·파손·탄소비용을 별도 계산합니다.")
    add_bullet(doc, "대게·한우·애플망고처럼 남은 시간이 짧은 품목은 먼 거리 배송보다 점포 인근 예약·당일 수령·현장 판매가 우선일 수 있습니다.")

    heading(doc, 2, "6.3 폐기·기부·벤더 회수")
    add_bullet(doc, "폐기 직전의 모든 상품을 할인 대상으로 만들지 않습니다. 식품은 소비기한·품질·보관조건을 기준으로 판매 가능 cutoff와 폐기 cutoff를 분리합니다.")
    add_bullet(doc, "폐기비는 운송·처리·증빙·분리배출 비용을 포함하고, 회계상 손상차손과 실제 현금비용을 분리합니다.")
    add_bullet(doc, "기부·회수는 브랜드·법규·세무·품질·수령처 capacity를 확인한 뒤 실행하며, ‘폐기비 0원’으로 단순 가정하지 않습니다.")
    add_bullet(doc, "포장재 자원순환(Project100·폐비닐) 효과는 상품 처분 수익과 별도 ESG KPI로 기록합니다.")
    add_source_ref(doc, "HDS-CIRCULAR")
    add_source_ref(doc, "ALLBARO")

    heading(doc, 2, "6.4 브랜드 이미지")
    add_table(doc, ["리스크", "나쁜 실행", "보호형 실행", "측정"], [
        ("가격 기준 하락", "정상가 상품을 반복 공개 타임세일", "대상·기간·수량 제한, 채널 분리", "정상가 판매 잠식·가격회복기간"),
        ("프리미엄 희소성 훼손", "명품/고가품 대량 할인", "예약·1:1·콘텐츠·승인된 별도 채널", "브랜드 규정 위반·고객 반응"),
        ("식품 신뢰", "기한 임박품을 무조건 할인", "품질 검사·표시·안전 문구·판매 cutoff", "클레임·반품·안전 사고"),
        ("ESG 그린워싱", "폐기량 감소만 홍보", "실제 처리증빙·재사용·자원순환 원장", "재활용률·증빙·고객 신뢰"),
    ], [1850, 2700, 3100, 2060], first_col_bold=True, font_size=8.5)

    heading(doc, 1, "7. 실운영 아키텍처·거버넌스 제안")
    heading(doc, 2, "7.1 권장 데이터·AI 계층")
    add_table(doc, ["계층", "책임", "비용 원장", "승인/감사"], [
        ("원천 연결", "ERP·POS·WMS·OMS·멤버십·정산·폐기", "추출·ETL·전송·저장", "스키마·동기화·접근권한"),
        ("정합성/규칙", "소유권·기한·가격·용량·법규 하드 차단", "규칙 실행·예외 처리", "규칙 버전·차단 사유"),
        ("예측/검색", "판매속도·탄력도·유사상품·정책·계약", "모델 추론·검색·embedding", "표본·신뢰구간·모델 버전"),
        ("전략 엔진", "기준선·후보·비용·하방·fallback", "시나리오 실행·수식", "입력 스냅샷·결과 해시"),
        ("설명 계층", "MD용 근거·메시지·요약", "LLM·토큰·캐시·도구", "프롬프트·출력·인용"),
        ("승인/실행", "담당자 승인·가격/재고 변경·결과 회수", "사람 시간·API·재시도", "승인자·시간·변경 전후값"),
        ("학습/감사", "예측오차·실제 가치·모델 개선", "평가·라벨링·재학습", "대조군·드리프트·롤백"),
    ], [1650, 3600, 2100, 2360], first_col_bold=True, font_size=8.4)
    add_callout(doc, "역할 분리", "규칙 엔진은 안전·소유권·기한을 결정하고, 수식/예측 모델은 수요·비용을 계산하며, 생성형 AI는 설명·요약·대화 보조를 담당합니다. 생성형 AI가 가격·재고를 직접 확정하지 않도록 API 권한을 분리합니다.", fill=GREEN_LIGHT, accent=GREEN)

    heading(doc, 2, "7.2 비용 예산·라우팅 정책")
    add_table(doc, ["구간", "대상", "허용 AI 비용 정책", "차단/승격"], [
        ("저위험 반복", "SAFE·CAUTION·기한 여유", "규칙·배치·캐시 우선, 설명은 템플릿", "새 데이터·예외만 모델 승격"),
        ("중위험", "WARNING·시즌 30~90일", "소형 예측+유사검색, 필요시 LLM 1회", "AI 원가/예상가치 비율 초과 시 담당자 확인"),
        ("고위험", "CRITICAL_NEAR·DEAD_STOCK", "고품질 모델/검색 허용하되 호출 횟수 상한", "안전·소유권·품질 누락 시 하드 차단"),
        ("긴급", "D-7 이하·식품", "규칙·품질검사·담당자 승인이 우선", "느린 LLM 대기 금지; 사후 설명 가능"),
        ("대규모", "그룹 예산·공동 프로모션", "본사 예외 검토와 비용분담 사전 합의", "점포·카테고리별 순가치가 음수면 공동 실행 금지"),
    ], [1750, 2700, 3200, 2060], first_col_bold=True, font_size=8.5)

    heading(doc, 2, "7.3 감사·데이터 보호")
    add_bullet(doc, "재고 전략 입력에는 고객 실명·연락처 대신 SKU·점포·집계 고객군·VIP eligibility flag 등 최소 데이터만 사용합니다.")
    add_bullet(doc, "고객 수준 데이터를 사용하는 경우 목적·법적 근거·보유기간·접근권한·위탁/국외이전·삭제를 별도 검토합니다.")
    add_bullet(doc, "LLM 요청에는 원가·계약·개인정보가 포함되지 않도록 사전 마스킹하고, 모델/프롬프트/데이터 해시를 로그에 남깁니다.")
    add_bullet(doc, "AI 결과에는 입력 스냅샷, 데이터 신선도, 모델 버전, 사용한 규칙·검색 근거, 사람 승인자, 실제 결과를 연결합니다.")
    add_bullet(doc, "개인정보위 생성형 AI 처리 기준의 전 과정 적법성·안전성·내부관리 원칙을 사전 검토합니다.")
    add_source_ref(doc, "PIPC-GENAI")

    heading(doc, 1, "8. 시뮬레이션 설계 보강안")
    heading(doc, 2, "8.1 현재 SimulationControls와 추가 입력")
    add_table(doc, ["구분", "현재 필드", "추가 필드 제안", "주의"], [
        ("판매 조건", "discountRate, couponRate, pointRate", "실제 할인부담주체·사용률·중복 규칙", "명목 혜택과 현금 비용을 분리"),
        ("채널", "onlineShareRate, freeShipping", "channel, pickupShare, deliveryCapacity", "온라인 60% 기본값은 가정"),
        ("운영", "packingCost, bundleAssemblyCost", "pickMinutes, QA cost, reworkRate", "점포별 활동기준원가"),
        ("반품", "returnRate", "returnState, resaleRate, refundTiming", "상품군별 정책 차이"),
        ("보관/폐기", "storageCostPerUnitDay, disposalCostPerUnit", "spaceValue, electricity, donationValue, vendorReturn", "회피비용·잔존가치 분리"),
        ("AI", "없음", "AI 원가 원장 전체, uncertainty, reserve", "최종 순가치에 차감"),
    ], [1500, 2600, 3500, 2110], first_col_bold=True, font_size=8.4)

    heading(doc, 2, "8.2 산출물 변경")
    add_table(doc, ["현재 결과", "추가 결과", "사용자에게 보여줄 문구"], [
        ("incrementalContribution", "netAfterAICost", "AI 비용 차감 후 기준선 대비 순가치"),
        ("confidenceScore", "confidenceInterval + dataQuality", "예상 범위·표본·신선도"),
        ("warningMessages", "hardStops + costWarnings + downside", "차단·비용·하방 경고를 분리"),
        ("costRows", "aiCostRows + executionCostRows", "AI 결정원가와 실행원가를 두 묶음으로 표시"),
        ("fallback", "fallbackBudget + approvalGate", "다음 단계의 최대 비용·승인 주체"),
    ], [2600, 3100, 4010], first_col_bold=True)
    add_internal_ref(doc, "src/lib/simulation.ts", "SimulationResult 확장 후보")

    heading(doc, 2, "8.3 검증용 테스트 시나리오")
    add_table(doc, ["테스트", "입력", "기대 결과", "실패 기준"], [
        ("식품 D-3", "소비기한·품질검사 누락", "추천이 아닌 하드 차단", "할인 전략이 후보로 노출"),
        ("특약매입", "purchaseType=특약매입·권한 미확인", "벤더 협의 대기", "직매입과 같은 할인 계산"),
        ("슬라이더 100회", "campaignDays/discount 변경", "수식만 재계산", "LLM/API 100회 호출"),
        ("AI 비용 초과", "C_AI_case > net value", "AI 추천 보류/저비용 경로 승격", "음수 순가치인데 승인 가능"),
        ("실제 오차", "실제 판매 86%, 기간 +3일", "varianceReason·fallback 학습", "예상값만 남고 원인 미수집"),
        ("브랜드 보호", "고가품 대량 할인", "가격정책 승인 필요", "자동 공개 타임세일 추천"),
    ], [1800, 3000, 2800, 2110], first_col_bold=True, font_size=8.5)

    heading(doc, 1, "9. KPI·운영 대시보드 정의")
    add_table(doc, ["KPI", "정의", "목표 방향", "원천"], [
        ("위험재고 처리율", "처리기한 내 판매·회수·기부·폐기 완료 수량/대상 수량", "상승", "POS/WMS/폐기"),
        ("위험조정 증분가치", "전략 결과−기준선−하방 reserve", "상승", "전략·정산·폐기"),
        ("AI 비용/케이스", "AI 원가 원장 합계/케이스", "하락 또는 가치 대비 통제", "AI usage·인프라·사람"),
        ("AI 비용/처리수량", "AI 원가/완료 처리수량", "하락", "AI 원장·실행"),
        ("예측오차", "예상 판매·기간·가치와 실제의 차이", "감소", "전략·실행"),
        ("fallback 전환율", "1차 전략에서 보정/처분 단계로 전환한 케이스 비율", "원인별 관리", "실행 로그"),
        ("브랜드 보호 지표", "정상가 잠식·클레임·가격회복기간", "하락", "POS·고객·브랜드"),
        ("식품 안전 지표", "기한 위반·클레임·리콜·검사 누락", "0", "품질·CS"),
        ("자원순환 지표", "폐기물 분리·재활용·재사용 증빙", "상승", "Allbaro·ESG"),
        ("승인 리드타임", "탐지→승인까지 시간", "위험등급별 단축", "workflow"),
    ], [1900, 4000, 1700, 2110], first_col_bold=True, font_size=8.5)

    heading(doc, 1, "10. 실행 로드맵과 데이터 요청 체크리스트")
    heading(doc, 2, "10.1 단계별 실행")
    add_table(doc, ["단계", "기간/범위", "구현·운영", "통과 기준"], [
        ("P0-정합성", "더현대 서울 4개 층 직매입", "원천 스키마·소유권·기한·기준선·AI 원가 원장", "데이터 품질·하드 차단 통과"),
        ("P1-전략", "패션·식품·리빙 3개 카테고리", "수식/예측·채널별 비용·fallback·승인", "실제 과거 케이스 백테스트"),
        ("P1-운영", "계열사 담당자 Pilot", "검토·승인·실행·실제 결과 회수", "예상 vs 실제 오차·순가치 검증"),
        ("P2-협업", "점포 이동·공동 프로모션", "비용분담·정산·그룹 거버넌스", "양쪽 계열사 순가치 양수"),
        ("P2-확장", "타 점포·타 계열사", "점포·카테고리별 비용·기한·법규 프로필", "공통 KPI와 독립 원장 동시 충족"),
    ], [1500, 2700, 3100, 2410], first_col_bold=True, font_size=8.5)

    heading(doc, 2, "10.2 첫 워크숍에서 확인할 질문")
    for q in [
        "직매입·특약매입·임대매장의 실제 소유권·할인권·반품권·비용분담은 계약상 어떻게 다릅니까?",
        "더현대 서울 2F·3F·B1·1F에서 위험재고를 판단할 때 실제로 쓰는 cutoff(보관일, 시즌 종료, 소비기한, 최소 판매속도)는 무엇입니까?",
        "식품의 소비기한·lot·품질검사·폐기 승인·기부/회수 가능 경로와 처리 증빙은 어디에 있습니까?",
        "온라인·근거리 배송·선물·방문수령·픽업의 재고 예약·취소·미수령·반품·재판매 프로세스는 무엇입니까?",
        "VIP/H.Point/앱/팝업/아울렛의 할인·포인트·정산·브랜드 적용 제외 조건은 어디에 정의되어 있습니까?",
        "케이스 1건의 평균·최대 MD 검토시간, 물류·재무·법무 협의시간은 얼마입니까?",
        "현대퓨처넷이 제공할 수 있는 ERP·메시징·사이니지·검색·모델 운영·보안 범위와 SLA는 무엇입니까?",
        "AI 추천으로 인해 판매가 늘었다고 인정할 귀속률과 대조군·A/B 또는 사후 비교 방식은 무엇입니까?",
        "AI 비용 예산은 계열사·점포·카테고리·위험등급 중 어떤 단위로 부담하고, 본사 공통비는 어떻게 배부합니까?",
        "가격·재고·폐기 실행은 어떤 시스템에서 누가 승인하고, 실패·중복·재시도·롤백을 어떻게 기록합니까?",
    ]:
        add_bullet(doc, q)

    heading(doc, 1, "11. 최종 권고 요약")
    add_callout(doc, "권고 1 — AI를 싸게 쓰는 것이 아니라, 비싼 판단에만 쓴다", "위험등급·기한·소유권·기본 수식은 규칙/통계로 처리하고, LLM은 전략 설명·예외·정책 문서 요약에 제한합니다. 슬라이더 조정은 재호출 없이 수식으로 처리합니다.", fill=GREEN_LIGHT, accent=GREEN)
    add_callout(doc, "권고 2 — 신선식품은 ‘가장 높은 이윤’보다 ‘가장 작은 하방’을 먼저", "D-7 이하 품목은 품질·안전·배송 capacity가 최우선입니다. 대게·애플망고·한우를 매일 재예측하고, 판매·기부·폐기의 cutoff를 분리하십시오.", fill="FFF8E7", accent=AMBER)
    add_callout(doc, "권고 3 — 패션·리빙은 가격이 아니라 노출과 채널을 단계화", "더현대 서울의 팝업·체류·큐레이션 강점을 활용하고, 공개 할인은 늦게·좁게·승인된 수량으로 사용합니다. 아울렛·별도 채널·벤더 협의를 fallback으로 둡니다.", fill=GREEN_LIGHT, accent=GREEN)
    add_callout(doc, "권고 4 — ‘AI 비용 차감 후 순가치’를 최우선 KPI로", "현재 프로젝트의 expectedNetContributionMargin만으로 승인하지 말고, AI·사람·운영·실패 reserve를 차감한 netAfterAICost와 기준선 대비 위험조정 가치를 함께 보여주십시오.", fill="FFF8E7", accent=AMBER)
    add_callout(doc, "권고 5 — 공개 사실과 내부 가정을 분리한 채 Pilot", "현대백화점 공개자료는 회사의 운영 철학·서비스 접점·ESG·AI/디지털 역량을 보여주지만 SKU별 재고 정책은 공개하지 않습니다. P0는 더현대 서울 직매입 4개 층의 과거 실제 케이스로 백테스트해야 합니다.", fill=GREEN_LIGHT, accent=GREEN)

    heading(doc, 1, "12. 조사 출처 및 근거 등록부")
    add_para(doc, "아래 출처는 2026-07-26 기준으로 확인한 공개 페이지입니다. 웹 페이지의 내용·가격·혜택·법규는 바뀔 수 있으므로 실운영 의사결정 전에 다시 확인하십시오.")
    table = doc.add_table(rows=1, cols=5)
    headers = ["코드", "출처", "일자", "지원 내용", "링크"]
    for i, h in enumerate(headers):
        set_cell_text(table.rows[0].cells[i], h, bold=True, color=DARK_BLUE, size=8.0)
        shade(table.rows[0].cells[i], LIGHT)
    for src in SOURCES:
        cells = table.add_row().cells
        set_cell_text(cells[0], src.code, bold=True, color=GREEN, size=7.7)
        set_cell_text(cells[1], src.title, size=7.7)
        set_cell_text(cells[2], src.date, size=7.7)
        set_cell_text(cells[3], src.supports, size=7.7)
        cells[4].text = ""
        p = cells[4].paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        add_hyperlink(p, "열기", src.url)
        for run in p.runs:
            set_run(run, size=7.7, color=BLUE)
    set_table_geometry(table, [1250, 2450, 1250, 3100, 1310], header_fill=LIGHT, font_size=7.7)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    heading(doc, 2, "프로젝트 내부 근거 파일")
    add_table(doc, ["파일", "문서에서 사용한 근거"], [
        ("docs/project-brief.md", "제품 목적·사용자·Anchor·범위·계열사 원칙"),
        ("docs/decision-policy.md", "기준선·증분 기여현금이익·비용 분류·하드 차단·AI/사람 책임"),
        ("src/lib/types.ts", "InventoryItem·StrategyOption·OptimizationCase 타입"),
        ("src/lib/mock-data.ts", "20개 재고·4개 층·위험등급·케이스·실행 결과·더미 비용"),
        ("src/lib/simulation.ts", "SimulationControls·SimulationResult·비용 계산·fallback"),
        ("src/app/strategy/generate/page.tsx", "4단계 AI 생성 목업 흐름"),
        ("src/app/strategy/[id]/simulate/simulate-client.tsx", "조정 패널·비용 브레이크다운·검토 요청"),
        ("src/app/strategy/execution/page.tsx", "AI 예상 vs 실제 결과·오차 회수"),
        ("src/components/inventory/product-detail-modal.tsx", "위험 분석·Do Nothing 기준선·가중치 목업"),
    ], [3200, 6160], first_col_bold=True, font_size=8.4)
    add_para(doc, "끝. 이 문서는 실제 가격 변경·재고 변경·폐기를 승인하지 않으며, 실운영 전 데이터·계약·법규·권한·예산 검증을 전제로 합니다.", italic=True, color=MUTED, size=9.5, after=0)
    add_header_footer(doc)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build_document()
