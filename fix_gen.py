import pathlib

content = '''# -*- coding: utf-8 -*-
import os
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

def create_document():
    doc = docx.Document()

    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    PRIMARY_GREEN = RGBColor(15, 76, 58)
    DARK_TEXT = RGBColor(30, 41, 59)
    GRAY_TEXT = RGBColor(100, 116, 139)

    normal_style = doc.styles['Normal']
    normal_style.font.name = 'Malgun Gothic'
    normal_style.font.size = Pt(10)
    normal_style.font.color.rgb = DARK_TEXT

    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_p.add_run("현대백화점 B2B AI 재고 수익 최적화 플랫폼\\n전체 기술 스택 및 아키텍처 종합 명세서")
    title_run.font.name = 'Malgun Gothic'
    title_run.font.size = Pt(18)
    title_run.font.bold = True
    title_run.font.color.rgb = PRIMARY_GREEN
    title_p.paragraph_format.space_after = Pt(15)

    sub_p = doc.add_paragraph()
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = sub_p.add_run("문서 버전: v1.0 | 작성일자: 2026-07-28 | 관제 대상: 더현대 서울 직매입 악성재고 관제 & 시뮬레이션 시스템")
    sub_run.font.name = 'Malgun Gothic'
    sub_run.font.size = Pt(9.5)
    sub_run.font.color.rgb = GRAY_TEXT
    sub_p.paragraph_format.space_after = Pt(25)

    def add_h1(text):
        h = doc.add_paragraph()
        h.paragraph_format.space_before = Pt(16)
        h.paragraph_format.space_after = Pt(6)
        h.paragraph_format.keep_with_next = True
        r = h.add_run(text)
        r.font.name = 'Malgun Gothic'
        r.font.size = Pt(13)
        r.font.bold = True
        r.font.color.rgb = PRIMARY_GREEN
        return h

    def add_h2(text):
        h = doc.add_paragraph()
        h.paragraph_format.space_before = Pt(12)
        h.paragraph_format.space_after = Pt(4)
        h.paragraph_format.keep_with_next = True
        r = h.add_run(text)
        r.font.name = 'Malgun Gothic'
        r.font.size = Pt(11)
        r.font.bold = True
        r.font.color.rgb = DARK_TEXT
        return h

    def add_p(text, bold_prefix=""):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.15
        if bold_prefix:
            r_b = p.add_run(bold_prefix)
            r_b.font.name = 'Malgun Gothic'
            r_b.font.bold = True
        r_t = p.add_run(text)
        r_t.font.name = 'Malgun Gothic'
        return p

    def set_cell_background(cell, fill_hex):
        shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
        cell._tc.get_or_add_tcPr().append(shd)

    def style_table(table, col_widths, headers, data):
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        hdr_cells = table.rows[0].cells
        for i, header_text in enumerate(headers):
            hdr_cells[i].text = header_text
            set_cell_background(hdr_cells[i], "0F4C3A")
            p = hdr_cells[i].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.font.name = 'Malgun Gothic'
                run.font.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)
                run.font.size = Pt(9)

        for row_idx, row_data in enumerate(data):
            row_cells = table.rows[row_idx + 1].cells
            bg_color = "F8FAFC" if row_idx % 2 == 1 else "FFFFFF"
            for col_idx, cell_value in enumerate(row_data):
                row_cells[col_idx].text = str(cell_value)
                set_cell_background(row_cells[col_idx], bg_color)
                p = row_cells[col_idx].paragraphs[0]
                p.paragraph_format.space_before = Pt(3)
                p.paragraph_format.space_after = Pt(3)
                p.paragraph_format.line_spacing = 1.15
                for run in p.runs:
                    run.font.name = 'Malgun Gothic'
                    run.font.size = Pt(8.5)

        for row in table.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Inches(w)

    add_h1("1. 프로젝트 개요 및 목적")
    add_p("본 시스템은 현대백화점 B2B 재고 의사결정 및 수익 최적화 운영 플랫폼입니다. 1차 대상을 더현대 서울 직매입 악성재고로 지정하여, 보관비·폐기비·반품률·잠식효과를 종합 고려한 '증분 기여현감이익(Incremental Cash Margin)'을 극대화합니다.")
    add_p("더현대 서울 재고 담당자가 일반 전략을 승인·실행하는 주체이며, 본사는 예외 리스크 및 그룹 공동 프로모션 조정 역할을 담당합니다.")

    add_h1("2. 영역별 기술 스택 분석 및 비교 (Trade-off Matrix)")

    add_h2("2.1 프론트엔드 (Frontend)")
    headers_fe = ["기술 스택", "필요성 & 개발 기능", "비교 대안 및 선택 이유", "트레이드오프 (Trade-off)"]
    widths_fe = [1.2, 2.3, 1.8, 1.2]
    data_fe = [
        ["React 19 + Vite", "B2B 오퍼레이션 타워 SPA 구축, 초고속 HMR 및 빌드", "vs Next.js (SSR)\\nB2B 내부 시스템으로 SEO가 필요 없으며 SSR 오버헤드 없이 빠른 빌드와 클라이언트 SPA 상태 유지가 뛰어남", "SEO 미지원 (B2B이므로 문제 없음)"],
        ["TanStack Query v5", "재고 API 캐싱, 상세진단 비동기 로딩, 전략생성 백그라운드 폴링", "vs Custom Fetch / Redux Thunk\\n서버 데이터의 자동 캐싱, 리프레시, 상태 관리를 선언적 코드로 처리", "클라이언트 UI 상태는 별도 전역 스토어 필요"],
        ["URL Search Params", "층(2F~1F), 카테고리, D-Day 필터링 상태의 URL 연동", "vs In-memory state\\n새로고침, 뒤로가기, URL 링크 공유 시에도 필터 조건이 100% 유지됨", "URL 길이가 다소 길어질 수 있음"],
        ["Zustand", "시뮬레이션 슬라이더 60fps 실시간 차트 연동, Multi-Select 장바구니, 전략 비교함", "vs Context API / Redux\\nContext API의 전체 불필요 리렌더링 문제를 Selector 기반 정밀 타겟 구독으로 완벽 해결", "스토어 구조 분리 관리 필요"],
        ["TanStack Table v8", "수천 개 직매입 SKU 스티키 헤더, 컬럼 정렬, 다중 선택, 페이징", "vs Basic Table / AG Grid\\n가볍고 Headless하여 디자인 자유도가 높으며 초고속 가상화(Virtualization) 지원", "shadcn/ui와 조합하여 직접 스타일링 필요"],
        ["Recharts", "할인율 vs 증분이익 vs 예상 판매량 듀얼 축 차트, 위험도 스태킹 바", "vs Chart.js / D3.js\\nReact 컴포넌트 친화적이며 SVG 기반으로 듀얼 축 및 실시간 데이터 업데이트가 매끄러움", "D3 대비 대량 캔버스 연산 제약 존재"],
        ["Tailwind CSS v4 + shadcn/ui", "현대백화점 그린(#0F4C3A) 디자인 시스템, 모달, 상태 배지", "vs Styled Components / MUI\\n빌드 타임 CSS 생성으로 무겁지 않고 프리셋 디자인 파편화가 없음", "HTML 클래스 유틸리티가 길어질 수 있음"],
        ["React Hook Form + Zod", "시뮬레이션 파라미터 조율 슬라이더 유효성 검증 및 폼 관리", "vs Formik / Plain State\\n비제어 컴포넌트 기반으로 폼 리렌더링 최소화 및 Zod 스키마 타입 안전성 제공", "Zod 스키마 정의 공수 필요"],
        ["Web Vitals", "LCP < 1.2s, INP < 100ms 등 UX 성능 정량 관제", "vs 측정 안 함\\n슬라이더 조율 및 화면 렌더링 품질을 정량 관제", "측정 스크립트 셋업 필요"],
        ["Vitest + RTL + Playwright", "단위 테스트, 컴포넌트 동작 검증 및 브라우저 E2E 자동화 테스트", "vs Jest / Cypress\\nVite 환경과 100% 호환되어 초고속 테스트 실행 및 E2E 시나리오 검증", "테스트 코드 작성 공수"]
    ]
    t_fe = doc.add_table(rows=len(data_fe) + 1, cols=4)
    style_table(t_fe, widths_fe, headers_fe, data_fe)

    add_h2("2.2 백엔드 (Backend)")
    headers_be = ["기술 스택", "필요성 & 개발 기능", "비교 대안 및 선택 이유", "트레이드오프 (Trade-off)"]
    widths_be = [1.2, 2.3, 1.8, 1.2]
    data_be = [
        ["Java 21 + Spring Boot 3.3", "RESTful API, 가상 스레드(Virtual Threads) 기반 동시 요청 고성능 처리", "vs Node.js / Python\\n대기업 엔터프라이즈 환경에서의 안정성, 멀티스레드 동시성, 풍부한 생태계", "Node.js 대비 초기 메모리 점유율 존재"],
        ["MyBatis 3.5+ (단독 채택, JPA 미사용)", "[단독 데이터 접근 계층] 재고 CRUD, 다중 조건 동적 필터링, 대용량 손실 집계 리포트, Window Function/CTE SQL 등 전체 DB 처리 전담", "vs JPA + QueryDSL / JPA 단독\\nJPA와 SQL 매퍼 동시 도입에 따른 복잡도를 배제하고, 개발자가 SQL을 100% 직접 작성 및 제어·튜닝 가능. 동적 쿼리는 MyBatis <if>, <choose>, <where> 태그 활용", "XML 매핑 파일 작성 및 쿼리 관리 필요 (자바 오타 컴파일 체크 미지원)"],
        ["Spring State Machine", "위험탐지➔검토➔본사승인➔실행➔완료 상태 전이 규칙 엄격 제어", "vs Enum + If/Else\\n상태 점프나 오작동을 차단하고 상태 변경 이벤트를 체계적으로 후킹함", "초기 트랜지션 구조 복잡도"],
        ["Spring Security + JWT", "더현대 서울 담당자 vs 본사 담당자 역할 기반 권한 제어 (RBAC)", "vs Session Auth\\n무상태(Stateless) API로 캐시 및 확장성이 우수함", "토큰 탈취 대비 무효화 레디스 연동 필요"],
        ["Spring Batch 5 + Quartz", "매일 새벽 02시 전 점포 직매입 재고 일일 손실액 및 위험도 자동 갱신", "vs Spring @Scheduled\\nQuartz DB 락(Clustering)을 통해 서버 이중화 시 중복 배치 실행 사고 완벽 방지", "Quartz DB 테이블 생성 및 설정 필요"],
        ["JUnit 5 + Mockito + Testcontainers", "서비스 로직 및 Docker 기반 실 PostgreSQL/Redis 실전 통합 테스트", "vs H2 메모리 DB\\nH2와 PostgreSQL 간 JSONB/SQL 차이 오류를 방지하고 실전 환경과 동일 검증", "Docker 기반 테스트 실행 시간 증가"]
    ]
    t_be = doc.add_table(rows=len(data_be) + 1, cols=4)
    style_table(t_be, widths_be, headers_be, data_be)

    add_h2("2.3 데이터베이스 및 캐시 (Database & Cache)")
    headers_db = ["기술 스택", "필요성 & 개발 기능", "비교 대안 및 선택 이유", "트레이드오프 (Trade-off)"]
    widths_db = [1.2, 2.3, 1.8, 1.2]
    data_db = [
        ["PostgreSQL 16+", "관계형 데이터 정규화 저장 + 시뮬레이션 파라미터 및 사후대처 트리 JSONB 저장", "vs MySQL 8.0 / Oracle 19c\\n100% 무료 오픈소스이면서 오라클급 CBO 최적화기 및 JSONB GIN 인덱싱 지원", "MySQL 대비 관리 숙련도 필요"],
        ["Redis 7.x", "실시간 위험재고 카운트 캐싱, 세션 관리, LLM 문장 생성 결과 Caching (0초 응답)", "vs In-memory HashMap\\n분산 환경에서 데이터 공유 가능 및 TTL 기반 자동 만료 지원", "메모리 관리 및 인메모리 휘발성 대비 필요"]
    ]
    t_db = doc.add_table(rows=len(data_db) + 1, cols=4)
    style_table(t_db, widths_db, headers_db, data_db)

    add_h2("2.4 AI 및 수리 최적화 마이크로서비스 (AI Service)")
    headers_ai = ["기술 스택", "필요성 & 개발 기능", "비교 대안 및 선택 이유", "트레이드오프 (Trade-off)"]
    widths_ai = [1.2, 2.3, 1.8, 1.2]
    data_ai = [
        ["Python 3.11+ + FastAPI", "수리 최적화 연산 및 ML 예측 전용 독립 마이크로서비스", "vs Spring AI 단독\\n파이썬 생태계의 머신러닝/수리 최적화 전용 라이브러리를 직접 활용 가능", "Spring Boot와의 HTTP 통신 제어 필요"],
        ["PuLP / SciPy", "보관비·폐기비·쿠폰 비용 반영 증분 기여현감이익 극대화 최적 할인율 계산", "vs 직접 작성 알고리즘\\n검증된 수리 방정식 최적화 엔진으로 오차 없는 정확한 마진 지점 산출", "방정식 제약조건 정의 필요"],
        ["scikit-learn / LightGBM", "과거 판매 이력 기반 수요 가격 탄력성 및 소진율 추정 ML 예측", "vs RAG / Vector DB\\n정형 숫자의 과거 데이터 학습에는 벡터 DB보다 ML 회귀 모델이 100% 정확함", "과거 데이터 품질 관리가 중요함"],
        ["Google Gemini 1.5 / Ollama / GPT-4o-mini", "현업 담당자용 정밀 진단 사유, 시나리오 추천 배경, 사후 대처 실행 가이드 문장 작성", "vs 로컬 LLM 파인튜닝\\nAPI 호출만으로 고품질 문장을 0원에 가깝게 활용 가능 (Gemini 하루 1,500회 무료)", "외부 네트워크 연동 필요 (Ollama 대체 가능)"]
    ]
    t_ai = doc.add_table(rows=len(data_ai) + 1, cols=4)
    style_table(t_ai, widths_ai, headers_ai, data_ai)

    add_h2("2.5 관제, 모니터링, Sentry 및 품질 검증 (Observability & Testing)")
    headers_obs = ["기술 스택", "필요성 & 개발 기능", "비교 대안 및 선택 이유", "트레이드오프 (Trade-off)"]
    widths_obs = [1.2, 2.3, 1.8, 1.2]
    data_obs = [
        ["Sentry", "프론트엔드/백엔드 실시간 런타임 에러 캡처 및 슬랙 알림", "vs 로그 파일 직접 확인\\n에러 발생 시 정확한 파일/라인 수, 콜스택, 사용자 동작 이력을 실시간 전달", "무료 플랜 트래픽 한도 (월 5,000건)"],
        ["Prometheus + Grafana", "서버 CPU, 메모리, DB 커넥션 풀, API RPS/Latency 실시간 시각화 관제", "vs CloudWatch 의존\\n무료 오픈소스로 가상 스레드 및 커넥션 풀 지표를 정밀 모니터링", "프로메테우스 에이전트 셋업 필요"],
        ["Loki + Promtail", "Logback JSON 로깅 수집 및 Grafana를 통한 중앙 로그 검색", "vs ELK Stack\\n엘라스틱서치 대비 메모리 사용량이 1/10 수준으로 매우 가볍고 직관적임", "전체 텍스트 색인 기능은 단순함"],
        ["OpenTelemetry + Jaeger", "React➔Spring Boot➔FastAPI➔PostgreSQL 구간 분산 트레이싱 추적", "vs 로그 분리 관찰\\n단일 TraceID로 서비스 간 병목 구간을 한눈에 식별 가능", "TraceID 전파 헤더 셋업 필요"],
        ["k6", "동시 접속자 500명 부하 상황에서의 P95 Latency 및 TPS 측정", "vs JMeter / Locust\\n자바스크립트 스크립트로 작성하기 쉬우며 리소스 소비가 극히 적음", "CLI 중심 조작 환경"]
    ]
    t_obs = doc.add_table(rows=len(data_obs) + 1, cols=4)
    style_table(t_obs, widths_obs, headers_obs, data_obs)

    add_h1("3. AI 예측 정확도 & 퀄리티 극대화 4대 최적화 전략")
    add_p("보관일수, 유통기한 D-Day, 과거 30/60/90일 판매 이동평균, 계절성, 할인율을 조합해 LightGBM 모델을 학습시킵니다. 할인율 상승 시 소진율 증가폭이 완만해지는 가격 탄력성 감쇄 곡선을 수학적으로 보정하여 현실적인 소진율을 추정합니다.", "1. ML 특성 공학 (Feature Engineering) & 탄력성 곡선 보정: ")
    add_p("과거 판매 이력이 부족한 신규 상품은 '동일 브랜드 ➔ 동일 카테고리 ➔ 동일 가격대' 유사 상품의 탄력성을 가중 평균하여 예측값을 보완하고, 표본 부족 시 담당자에게 '신뢰도 65% 경고'를 표시합니다.", "2. 신규/비인기 SKU 콜드스타트(Cold-Start) 방지: ")
    add_p("LLM 호출 시 지정된 JSON 포맷만 반환하도록 강제하여 파싱 에러를 차단하며, 백화점 오퍼레이션 격식 템플릿을 적용합니다. 동일 재고 조건의 설명 문장은 Redis에 캐싱하여 0초 응답을 보장합니다.", "3. LLM 구조화 출력(Structured Output) & Redis 0초 캐싱: ")
    add_p("실제 실행된 프로모션 결과(실제 판매량, 증분이익)를 DB로 회수하여 AI 예상치와의 오차(MAPE)를 기록하고, 매월 파이썬 ML 모델이 스스로 재학습하여 정밀도를 지속적으로 향상시킵니다.", "4. 피드백 재학습 루프 (Closed-Loop Feedback System): ")

    add_h1("4. 종합 시스템 아키텍처 다이어그램")
    add_p("전체 시스템은 Frontend (React SPA) ➔ Reverse Proxy (Nginx) ➔ Core Backend (Spring Boot) ➔ AI Microservice (Python FastAPI) ➔ Database (PostgreSQL/Redis) ➔ Observability (Prometheus/Loki/Sentry)로 유기적으로 연동됩니다.")

    os.makedirs("/Users/junha/coding/stock/deliverables", exist_ok=True)
    doc_path = "/Users/junha/coding/stock/deliverables/hyundai_inventory_tech_stack_architecture.docx"
    doc.save(doc_path)
    print(f"Docx saved to {doc_path}")

if __name__ == "__main__":
    create_document()
'''

pathlib.Path('generate_docs.py').write_text(content, encoding='utf-8')
print("generate_docs.py fixed successfully")
