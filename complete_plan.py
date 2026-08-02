import docx
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import qn, nsdecls

def set_cell_background(cell, color_hex):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)

def set_cell_margins(cell, top=120, bottom=120, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = parse_xml(f'<w:tcMar {nsdecls("w")}><w:top w:w="{top}" w:type="dxa"/><w:bottom w:w="{bottom}" w:type="dxa"/><w:left w:w="{left}" w:type="dxa"/><w:right w:w="{right}" w:type="dxa"/></w:tcMar>')
    tcPr.append(tcMar)

def set_table_borders(table, color="D3D3D3", sz="4", val="single"):
    tblPr = table._tbl.tblPr
    borders = parse_xml(f'''
        <w:tblBorders {nsdecls("w")}>
            <w:top w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>
            <w:bottom w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>
            <w:insideH w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>
            <w:insideV w:val="none"/>
            <w:left w:val="none"/>
            <w:right w:val="none"/>
        </w:tblBorders>
    ''')
    tblPr.append(borders)

def style_heading_paragraph(p, text, size_pt=12, bold=True, space_before=12, space_after=6, color_rgb=(0, 51, 102)):
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = 1.2
    p.paragraph_format.keep_with_next = True
    p.text = ""
    run = p.add_run(text)
    run.font.name = "Malgun Gothic"
    run.font.size = Pt(size_pt)
    run.bold = bold
    run.font.color.rgb = RGBColor(*color_rgb)

def add_body_p(doc, text, font_name="Malgun Gothic", size_pt=10, bold=False, space_before=2, space_after=4, color_rgb=(51,51,51), align=WD_ALIGN_PARAGRAPH.LEFT, line_spacing=1.15):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = line_spacing
    run = p.add_run(text)
    run.font.name = font_name
    run.font.size = Pt(size_pt)
    run.bold = bold
    run.font.color.rgb = RGBColor(*color_rgb)
    return p

doc_path = '/Users/junha/Downloads/프로젝트 수행계획서_양식.docx'
doc = docx.Document(doc_path)

print("Script template ready")
