import os
import docx
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import qn, nsdecls

def set_cell_background(cell, color_hex):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)

def set_cell_margins(cell, top=120, bottom=120, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = parse_xml(f'<w:tcMar {nsdecls("w")}><w:top w:w="{top}" w:type="dxa"/><w:bottom w:w="{bottom}" w:type="dxa"/><w:left w:w="{left}" w:type="dxa"/><w:right w:w="{right}" w:type="dxa"/></w:tcMar>')
    tcPr.append(tcMar)

def set_table_borders(table, color="D3D3D3", sz="4", val="single"):
    tblPr = table._tbl.tblPr
    borders = parse_xml(f'''
        <w:tblBorders {nsdecls("w")}>
            <w:top w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>
            <w:bottom w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>
            <w:insideH w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>
            <w:insideV w:val="none"/>
            <w:left w:val="none"/>
            <w:right w:val="none"/>
        </w:tblBorders>
    ''')
    tblPr.append(borders)

def format_cell_text(cell, text, font_name="Malgun Gothic", size_pt=9.5, bold=False, color_rgb=(51,51,51), align=WD_ALIGN_PARAGRAPH.LEFT, space_before=3, space_after=3):
    p = cell.paragraphs[0]
    p.text = ""
    p.alignment = align
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = 1.15
    run = p.add_run(text)
    run.font.name = font_name
    run.font.size = Pt(size_pt)
    run.bold = bold
    run.font.color.rgb = RGBColor(*color_rgb)

doc_path = '/Users/junha/Downloads/프로젝트 수행계획서_양식.docx'
doc = docx.Document(doc_path)

# Clear existing body elements to rebuild cleanly with template's tables & styles
body = doc._body._element
for child in list(body):
    if child.tag.endswith(('p', 'tbl')):
        body.remove(child)

def add_p(text="", size_pt=10, bold=False, space_before=3, space_after=4, color_rgb=(51,51,51), align=WD_ALIGN_PARAGRAPH.LEFT, line_spacing=1.15, font_name="Malgun Gothic"):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = line_spacing
    if text:
        run = p.add_run(text)
        run.font.name = font_name
        run.font.size = Pt(size_pt)
        run.bold = bold
        run.font.color.rgb = RGBColor(*color_rgb)
    return p

def add_sec_h(text):
    return add_p(text, size_pt=13, bold=True, space_before=14, space_after=6, color_rgb=(0, 51, 102))

def add_sub_h(text):
    return add_p(text, size_pt=11, bold=True, space_before=8, space_after=4, color_rgb=(0, 51, 102))

# 1. Document Title
add_p("프 로 젝 트 수 행 계 획 서", size_pt=22, bold=True, space_before=12, space_after=18, color_rgb=(0, 51, 102), align=WD_ALIGN_PARAGRAPH.CENTER)

# 2. Table 1 (Header Info Table)
t1 = doc.add_table(rows=4, cols=2)
set_table_borders(t1, color="B0C4DE")
t1_data = [
    ("팀명", "[2팀] HEYDI Stock (현대백화점 B2B AI 재고 최적화 팀)"),
    ("프로젝트명(주제)", "현대백화점 B2B AI 재고 수익 최적화 플랫폼 (Heydi Stock Pro)"),
    ("멘토", "김영만 선임 (현대백화점그룹 본사 재고전략팀)"),
    ("멘티(교육생)", "김준하 (풀스택 개발자)")
]
for idx, (lbl, val) in enumerate(t1_data):
    format_cell_text(t1.rows[idx].cells[0], lbl, size_pt=9.5, bold=True, color_rgb=(0, 51, 102), align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_background(t1.rows[idx].cells[0], "F0F4F8")
    set_cell_margins(t1.rows[idx].cells[0])
    format_cell_text(t1.rows[idx].cells[1], val, size_pt=9.5, bold=False, color_rgb=(51,51,51))
    set_cell_margins(t1.rows[idx].cells[1])

add_p()

# 3. Section I: Project Overview
add_sec_h("I. 프로젝트 개요")
t2 = doc.add_table(rows=8, cols=2)
set_table_borders(t2, color="B0C4DE")
t2_data = [
    ("프로젝트명", "현대백화점 B2B AI 재고 수익 최적화 플랫폼 (Heydi Stock Pro)"),
    ("프로젝트 기간", "2026.07.20 ~ 2026.08.15 (4주)"),
    ("목적", "현대백화점 직매입 재고(더현대 서울 2F/3F/B1/1F)의 장기보관 및 유통기한 임박 손실을 새벽 자동 탐지(Spring Batch + Quartz)하고, 수리 최적화 연산기(Python SciPy/PuLP)와 생성형 AI(OpenAI GPT-4o-mini)를 결합하여 Do Nothing Baseline 대비 증분 기여현금이익을 극대화하는 B2B 의사결정 지원 플랫폼 구축"),
    ("키워드", "AI 재고 최적화, 기여현금이익(Contribution Cash Profit), Do Nothing Baseline, 수리 제약 최적화, B2B 오퍼레이션 타워, Oracle Retail RMS"),
    ("핵심기술", "수리 제약 최적화 알고리즘 (SciPy/PuLP), 생성형 AI 판단 근거 연동 (OpenAI GPT-4o-mini API), 대용량 위험재고 자동 탐지 파이프라인 (Spring Batch 5 + Quartz)"),
    ("사용프로그램", "React 18/19, Vite, TypeScript, Zustand, TanStack Table/Query, Recharts, Java 21, Spring Boot 3.3, QueryDSL, Python 3.11+, FastAPI, PostgreSQL 16, Redis, OpenAI API"),
    ("참고문헌", "Character Region Awareness for Text Detection (CVPR 2019), Oracle Retail Merchandising System (RMS) In-Season Inventory Analytics Guide (2024), 현대백화점 지속가능경영보고서 (2024)"),
    ("기대효과", "일일 보관/폐기 손실 사전 회피, Do Nothing 대비 증분 기여현금이익 15% 이상 증대, 5대 하드 차단 제약조건 자동제어로 브랜드 가치 보호, 재고 최적화 의사결정 소요 시간 80% 단축")
]
for idx, (lbl, val) in enumerate(t2_data):
    format_cell_text(t2.rows[idx].cells[0], f"  {lbl}", size_pt=9.5, bold=True, color_rgb=(0, 51, 102))
    set_cell_background(t2.rows[idx].cells[0], "F0F4F8")
    set_cell_margins(t2.rows[idx].cells[0])
    format_cell_text(t2.rows[idx].cells[1], val, size_pt=9.5, bold=False, color_rgb=(51,51,51))
    set_cell_margins(t2.rows[idx].cells[1])

add_p()

# 4. Section II: Project Scope
add_sec_h("II. 프로젝트 범위")
add_sub_h(" 1. 작업명세서 ")

t3 = doc.add_table(rows=8, cols=2)
set_table_borders(t3, color="B0C4DE")
t3_data = [
    ("업무", "업무범위"),
    ("데이터 전처리 및 DB 설계", "더현대 서울 직매입 재고 Master, 보관비/원가/유통기한/시즌 D-Day, 승인이력 PostgreSQL 16 스키마 정규화 설계 및 Mock 데이터 구축"),
    ("백엔드 Core 및 동적 검색 개발", "Spring Boot 3.3, QueryDSL 기반 층/카테고리/D-Day/위험도 다중조건 필터링 API 및 가상 스레드(Virtual Threads) 기반 동시 시뮬레이션 연동"),
    ("새벽 대용량 탐지 배치 개발", "Spring Batch 5 + Quartz Scheduler 기반 매일 02시 전점 재고 일일 보관 손실액 계산 및 위험도 5단계(SAFE~DEAD) 자동 분류 파이프라인 개발"),
    ("수리 최적화 연산 엔진 개발", "Python 3.11 FastAPI 기반 SciPy/PuLP 마진 연산기, 할인율 반응 곡선, 번들 매칭 및 60fps 실시간 슬라이더 시뮬레이션 엔드포인트 구축"),
    ("생성형 AI 판단 근거 연동", "OpenAI GPT-4o-mini API 연동을 통한 위험 사유 텍스트 생성, A/B/C 시나리오 추천 배경 요약, 사후 대처 트리(Fallback Action Plan) 지침 자동 작성"),
    ("프론트엔드 B2B UI 개발", "React 18/19, Zustand, TanStack Table/Query, Recharts 기반 통합 관제, 시뮬레이션 워크벤치, 성과 비교 관제 화면 개발"),
    ("검증 및 최종 보고", "E2E 통합 시나리오 테스트, Oracle Retail 지표 검증, 최종 수행계획서 작성 및 사용자 설명 웹사이트(explainer-site) 구축·배포")
]
for idx, (lbl, val) in enumerate(t3_data):
    bg_color = "D9E1F2" if idx == 0 else "FFFFFF"
    text_color = (0, 51, 102) if idx == 0 else (51, 51, 51)
    bold = True if idx == 0 else False
    align = WD_ALIGN_PARAGRAPH.CENTER if idx == 0 else WD_ALIGN_PARAGRAPH.LEFT
    
    format_cell_text(t3.rows[idx].cells[0], f" {lbl}", size_pt=9.5, bold=bold, color_rgb=text_color, align=align)
    if bg_color != "FFFFFF": set_cell_background(t3.rows[idx].cells[0], bg_color)
    set_cell_margins(t3.rows[idx].cells[0])
    
    format_cell_text(t3.rows[idx].cells[1], val, size_pt=9.5, bold=bold, color_rgb=text_color, align=align)
    if bg_color != "FFFFFF": set_cell_background(t3.rows[idx].cells[1], bg_color)
    set_cell_margins(t3.rows[idx].cells[1])

add_p()
add_sub_h(" 2. 방법론 및 접근법 ")
add_p("* 방법론 : 빅데이터 분석 및 Agile 반복 개발 방법론을 채택하여 분석기획, 도메인/데이터 설계, AI/연산 엔진 및 백엔드 구축, 프론트엔드 통합, 검증 및 전개 단계로 순차적·반복적 개발 프로젝트 진행", size_pt=9.5)
add_p("* 접근법 : 하향식 접근법(Top-down)을 채택하여 현황 분석을 통해 도출된 과제(현대백화점 직매입 재고 손실 최소화)에 대하여 프로젝트를 진행하며, 수리 방정식 100% 정밀 연산과 LLM 자연어 지침 생성을 결합한 하이브리드 AI 접근법 적용", size_pt=9.5, space_after=8)

# 5. Section III: Infrastructure & Background
add_sec_h("Ⅲ. 기반구축 내용 ")
add_sub_h(" 1. 기술개발배경 ")
add_p("* 본 프로젝트는 현대백화점 직매입 재고의 시즌 경과 및 소비기한 임박으로 발생하는 무분별한 덤핑/폐기 손실을 방지하고, 당면한 재고 처리 의사결정 효율성을 극대화하는 것을 일차적인 목표로 함.", size_pt=9.5)
add_p("* 기존 커머스 및 ERP/재고관리 시스템은 단순 일률적 할인이나 단순 폐기에 의존하여 보관비·폐기비·반품비·브랜드 훼손·정상판매 잠식비용을 종합적으로 반영하지 못함.", size_pt=9.5)
add_p("* 따라서 방치 시 기준선(Do Nothing Baseline) 대비 '위험조정 증분 기여현금이익'을 산출하고, 담당자가 직접 조율·검토·승인할 수 있는 B2B AI 오퍼레이션 타워 구축이 필수적임.", size_pt=9.5, space_after=8)

add_sub_h(" 2. 소프트웨어 구성도 ")
add_p("[ Client: React 18/19 SPA ] <──REST API──> [ Backend: Spring Boot 3.3 / QueryDSL ] <──HTTP──> [ AI Engine: Python FastAPI (SciPy/PuLP) + OpenAI GPT API ]\n                                              └───────── Database: PostgreSQL 16 & Redis ───────┘\n                                              └───────── Batch: Spring Batch 5 + Quartz Scheduler ─┘", font_name="Courier New", size_pt=8.5, color_rgb=(0, 51, 102), space_after=12)

# 6. Section IV: Implementation & Team Roles
add_sec_h("Ⅳ. 프로젝트 추진체계")
add_sub_h(" 1. 프로젝트 수행조직도 ")
add_p("                    [ 프로젝트 수행 총괄 (김준하) ]\n                                   │\n       ┌───────────────────────────┼───────────────────────────┐\n       ▼                           ▼                           ▼\n[ 풀스택 & AI 개발 ]     [ 도메인 & 전략 멘토 ]     [ 오퍼레이션 검증 ]\n (김준하 / 개발총괄)      (김영만 / 수석 MD)        (문주성 / 책임 MD)", font_name="Courier New", size_pt=8.5, color_rgb=(0, 51, 102), space_after=10)

add_sub_h(" 2. 구성원 역할")
t4 = doc.add_table(rows=4, cols=2)
set_table_borders(t4, color="B0C4DE")
t4_data = [
    ("구분", "역할 및 책임"),
    ("김준하", "프로젝트 총괄 / 백엔드(Spring Boot 3.3, Spring Batch), 파이썬 수리 최적화 엔진, 프론트엔드 B2B UI 개발, DB 스키마 설계 및 통합 배포"),
    ("김영만", "도메인 멘토 / 현대백화점 직매입 재고 의사결정 정책 수립, Oracle Retail RMS 분석 지표 정의, 증분 기여현금이익 가드레일 검토"),
    ("문주성", "오퍼레이션 검증 / 더현대 서울 패션 부문 아울렛 이관 권한 및 브랜드 할인 상한 5대 하드 차단 제약조건 검증")
]
for idx, (lbl, val) in enumerate(t4_data):
    bg_color = "D9E1F2" if idx == 0 else "FFFFFF"
    text_color = (0, 51, 102) if idx == 0 else (51, 51, 51)
    bold = True if idx == 0 else False
    align = WD_ALIGN_PARAGRAPH.CENTER if (idx == 0 or idx > 0 and len(lbl) < 5) else WD_ALIGN_PARAGRAPH.LEFT
    
    format_cell_text(t4.rows[idx].cells[0], f" {lbl}", size_pt=9.5, bold=bold, color_rgb=text_color, align=align)
    if bg_color != "FFFFFF": set_cell_background(t4.rows[idx].cells[0], bg_color)
    set_cell_margins(t4.rows[idx].cells[0])
    
    format_cell_text(t4.rows[idx].cells[1], val, size_pt=9.5, bold=bold, color_rgb=text_color, align=WD_ALIGN_PARAGRAPH.LEFT)
    if bg_color != "FFFFFF": set_cell_background(t4.rows[idx].cells[1], bg_color)
    set_cell_margins(t4.rows[idx].cells[1])

add_p()

# 7. Section V: Management Process
add_sec_h("Ⅴ. 관리 프로세스 계획")
add_sub_h(" 1. 세부일정 추진 계획 ")
t_sch = doc.add_table(rows=7, cols=5)
set_table_borders(t_sch, color="B0C4DE")

sch_headers = ["세부개발내용", "1주 (7/20~7/26)", "2주 (7/27~8/02)", "3주 (8/03~8/09)", "4주 (8/10~8/15)"]
for c_idx, text in enumerate(sch_headers):
    format_cell_text(t_sch.rows[0].cells[c_idx], text, size_pt=9.0, bold=True, color_rgb=(0, 51, 102), align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_background(t_sch.rows[0].cells[c_idx], "D9E1F2")
    set_cell_margins(t_sch.rows[0].cells[c_idx])

sch_rows = [
    ("요구사항 정의 & DB 스키마 설계", "■■■■■", "", "", ""),
    ("백엔드 API & QueryDSL 동적검색 개발", "", "■■■■■", "", ""),
    ("Python 수리 최적화 마진 연산기 개발", "", "■■■■■", "", ""),
    ("Spring Batch 새벽 탐지 & OpenAI API 연동", "", "", "■■■■■", ""),
    ("프론트엔드 B2B UI & 시뮬레이션 워크벤치", "", "", "■■■■■", ""),
    ("E2E 시나리오 테스트 & 산출물/웹사이트 구축", "", "", "", "■■■■■")
]

for r_idx, r_data in enumerate(sch_rows, start=1):
    for c_idx, val in enumerate(r_data):
        align = WD_ALIGN_PARAGRAPH.LEFT if c_idx == 0 else WD_ALIGN_PARAGRAPH.CENTER
        color = (0, 102, 204) if "■" in val else (51, 51, 51)
        bold = True if "■" in val else False
        format_cell_text(t_sch.rows[r_idx].cells[c_idx], val, size_pt=8.5, bold=bold, color_rgb=color, align=align)
        set_cell_margins(t_sch.rows[r_idx].cells[c_idx])

add_p()
add_sub_h(" 2. 단계별 산출물 ")
t_out = doc.add_table(rows=7, cols=4)
set_table_borders(t_out, color="B0C4DE")

out_headers = ["단계", "산출물", "완료일", "비고"]
for c_idx, text in enumerate(out_headers):
    format_cell_text(t_out.rows[0].cells[c_idx], text, size_pt=9.0, bold=True, color_rgb=(0, 51, 102), align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_background(t_out.rows[0].cells[c_idx], "D9E1F2")
    set_cell_margins(t_out.rows[0].cells[c_idx])

out_rows = [
    ("계획수립", "프로젝트 수행계획서", "07/22", "본 문서"),
    ("분석", "요구사항정의서 & 의사결정 정책서", "07/24", "Do Nothing 기준선 정의"),
    ("설계", "DB 스키마 설계서 & AI 최적화 수식 명세서", "07/27", "PostgreSQL / PuLP 연산"),
    ("개발", "백엔드 API & Python AI 엔진 & React B2B 프론트엔드", "08/08", "GitHub 소스코드 반영"),
    ("테스트", "E2E 시나리오 테스트결과서 & 사용설명 웹사이트", "08/12", "explainer-site 배포"),
    ("완료보고", "최종 프로젝트 완료보고서", "08/15", "발표자료 및 완료보고서")
]

for r_idx, r_data in enumerate(out_rows, start=1):
    for c_idx, val in enumerate(r_data):
        align = WD_ALIGN_PARAGRAPH.CENTER if c_idx in [0, 2] else WD_ALIGN_PARAGRAPH.LEFT
        format_cell_text(t_out.rows[r_idx].cells[c_idx], val, size_pt=8.5, bold=False, color_rgb=(51,51,51), align=align)
        set_cell_margins(t_out.rows[r_idx].cells[c_idx])

add_p()

# 8. Section VI: Significance & Importance
add_sec_h("Ⅵ. 개발기술의 의의 및 중요성")
add_p("* 개발 기술의 의의", size_pt=10, bold=True, color_rgb=(0, 51, 102))
add_p("확률적 LLM의 환각(Hallucination) 위험을 원천 차단하기 위해 '수리 제약 최적화 연산기(100% 수리적 정밀성)'와 '생성형 LLM(자연어 지침 가이드)'을 완벽히 분리 및 결합한 하이브리드 AI 아키텍처를 제시함.", size_pt=9.5)
add_p("백엔드(Spring Boot 3.3)와 AI 엔진(Python FastAPI)을 MSA 구조로 독립 분리하여, 수리 최적화 알고리즘 수정 시에도 메인 시스템 재부팅 없이 유연한 확장 및 배포가 가능함.", size_pt=9.5, space_after=8)

add_p("* 중요성", size_pt=10, bold=True, color_rgb=(0, 51, 102))
add_p("현대백화점 패션/식품 직매입 재고에 Oracle Retail 기준 5대 분석지표(ROS, Sell Thru %, WOS, Lift %, STD/BTA)와 5대 하드 차단 제약조건(식품안전, 소유권, 물류용량, 브랜드상한, AI ROI)을 최초로 정밀 적용하여 실무 현장 도입 가능성을 극대화함.", size_pt=9.5, space_after=12)

# 9. Section VII: Application & Expected Effects
add_sec_h("Ⅶ 활용방안/적용분야 및 기대효과 ")
add_p("* 활용방안", size_pt=10, bold=True, color_rgb=(0, 51, 102))
add_p("더현대 서울 4개 층(2F 여성패션, 3F 남성/잡화, B1 식품관, 1F 뷰티/리빙) 직매입 재고 관제 적용을 시작으로 현대백화점 15개 점포, 9개 현대아울렛, 현대식품관/Hmall 등 전사 유통 망으로 확산 가능.", size_pt=9.5)
add_p("직매입 수치 연산 체계를 바탕으로 입점 브랜드 특약매입 재고 대상 제안형 판촉 및 프로모션 시뮬레이션 솔루션으로 연계 활용 가능.", size_pt=9.5, space_after=8)

add_p("* 기대효과", size_pt=10, bold=True, color_rgb=(0, 51, 102))
add_p("악성 재고 보관/폐기 손실액 25% 이상 감축, 아무것도 하지 않는 방치(Do Nothing) 대비 증분 기여현금이익 15% 이상 증대.", size_pt=9.5)
add_p("5대 하드 차단 제약조건 자동 검증을 통한 브랜드 가치 보호 및 법규 준수 보장, 재고 최적화 의사결정 소요 시간 80% 단축.", size_pt=9.5, space_after=12)

# Save clean docx
output_path = '/Users/junha/coding/stock/deliverables/프로젝트 수행계획서_양식.docx'
doc.save(output_path)

print(f"Clean document successfully written and saved to {output_path}")
