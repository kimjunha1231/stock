import docx
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import qn, nsdecls

def set_cell_background(cell, color_hex):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)

def set_cell_margins(cell, top=120, bottom=120, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = parse_xml(f'<w:tcMar {nsdecls("w")}><w:top w:w="{top}" w:type="dxa"/><w:bottom w:w="{bottom}" w:type="dxa"/><w:left w:w="{left}" w:type="dxa"/><w:right w:w="{right}" w:type="dxa"/></w:tcMar>')
    tcPr.append(tcMar)

def set_table_borders(table, color="D3D3D3", sz="4", val="single"):
    tblPr = table._tbl.tblPr
    borders = parse_xml(f'''
        <w:tblBorders {nsdecls("w")}>
            <w:top w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>
            <w:bottom w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>
            <w:insideH w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>
            <w:insideV w:val="none"/>
            <w:left w:val="none"/>
            <w:right w:val="none"/>
        </w:tblBorders>
    ''')
    tblPr.append(borders)

doc_path = '/Users/junha/Downloads/프로젝트 수행계획서_양식.docx'
doc = docx.Document(doc_path)

print(f"Original DocumentLoaded: {len(doc.paragraphs)} paragraphs, {len(doc.tables)} tables")

# Table 1: Key Info Table
t1 = doc.tables[0]
t1.rows[0].cells[1].paragraphs[0].text = "[2팀] HEYDI Stock (현대백화점 B2B AI 재고 최적화 팀)"
t1.rows[1].cells[1].paragraphs[0].text = "현대백화점 B2B AI 재고 수익 최적화 플랫폼 (Heydi Stock Pro)"
t1.rows[2].cells[1].paragraphs[0].text = "김영만 선임 (현대백화점그룹 본사 재고전략팀)"
t1.rows[3].cells[1].paragraphs[0].text = "김준하 (풀스택 개발자)"

for r in t1.rows:
    for c in r.cells:
        for p in c.paragraphs:
            p.paragraph_format.space_before = Pt(3)
            p.paragraph_format.space_after = Pt(3)
            for run in p.runs:
                run.font.name = "Malgun Gothic"
                run.font.size = Pt(10)

# Table 2: Project Overview Table
t2 = doc.tables[1]
t2_data = [
    ("프로젝트명", "현대백화점 B2B AI 재고 수익 최적화 플랫폼 (Heydi Stock Pro)"),
    ("프로젝트 기간", "2026.07.20 ~ 2026.08.15 (4주)"),
    ("목적", "현대백화점 직매입 재고(더현대 서울 2F/3F/B1/1F)의 장기보관 및 유통기한 임박 손실을 새벽 자동 탐지(Spring Batch + Quartz)하고, 수리 최적화 연산기(Python SciPy/PuLP)와 생성형 AI(OpenAI GPT-4o-mini)를 결합하여 Do Nothing Baseline 대비 증분 기여현금이익을 극대화하는 B2B 의사결정 지원 플랫폼 구축"),
    ("키워드", "AI 재고 최적화, 기여현금이익(Contribution Cash Profit), Do Nothing Baseline, 수리 제약 최적화, B2B 오퍼레이션 타워, Oracle Retail RMS"),
    ("핵심기술", "수리 제약 최적화 알고리즘 (SciPy/PuLP), 생성형 AI 판단 근거 연동 (OpenAI GPT-4o-mini API), 대용량 위험재고 자동 탐지 파이프라인 (Spring Batch 5 + Quartz)"),
    ("사용프로그램", "React 18/19, Vite, TypeScript, Zustand, TanStack Table/Query, Recharts, Java 21, Spring Boot 3.3, QueryDSL, Python 3.11+, FastAPI, PostgreSQL 16, Redis, OpenAI API"),
    ("기대효과", "일일 보관/폐기 손실 사전 회피, Do Nothing 대비 증분 기여현금이익 15% 이상 증대, 5대 하드 차단 제약조건 자동제어로 브랜드 가치 보호, 재고 최적화 의사결정 소요 시간 80% 단축")
]

for idx, (label, val) in enumerate(t2_data):
    if idx < len(t2.rows):
        row = t2.rows[idx]
    else:
        row = t2.add_row()
    row.cells[0].paragraphs[0].text = f"  {label}"
    row.cells[1].paragraphs[0].text = val

# Add extra row for 참고문헌 if desired to match example PDF
row_ref = t2.add_row()
row_ref.cells[0].paragraphs[0].text = "  참고문헌"
row_ref.cells[1].paragraphs[0].text = "Character Region Awareness for Text Detection (CVPR 2019), Oracle Retail Merchandising System (RMS) In-Season Inventory Analytics Guide (2024), 현대백화점 지속가능경영보고서 (2024)"

for r in t2.rows:
    for c in r.cells:
        for p in c.paragraphs:
            p.paragraph_format.space_before = Pt(3)
            p.paragraph_format.space_after = Pt(3)
            p.paragraph_format.line_spacing = 1.15
            for run in p.runs:
                run.font.name = "Malgun Gothic"
                run.font.size = Pt(9.5)

# Table 3: Work Breakdown Table
t3 = doc.tables[2]
t3_data = [
    ("데이터 전처리 및 DB 설계", "더현대 서울 직매입 재고 Master, 보관비/원가/유통기한/시즌 D-Day, 승인이력 PostgreSQL 16 스키마 정규화 설계 및 Mock 데이터 구축"),
    ("백엔드 Core 및 동적 검색 개발", "Spring Boot 3.3, QueryDSL 기반 층/카테고리/D-Day/위험도 다중조건 필터링 API 및 가상 스레드(Virtual Threads) 기반 동시 시뮬레이션 연동"),
    ("새벽 대용량 탐지 배치 개발", "Spring Batch 5 + Quartz Scheduler 기반 매일 02시 전점 재고 일일 보관 손실액 계산 및 위험도 5단계(SAFE~DEAD) 자동 분류 파이프라인 개발"),
    ("수리 최적화 연산 엔진 개발", "Python 3.11 FastAPI 기반 SciPy/PuLP 마진 연산기, 할인율 반응 곡선, 번들 매칭 및 60fps 실시간 슬라이더 시뮬레이션 엔드포인트 구축"),
    ("생성형 AI 판단 근거 연동", "OpenAI GPT-4o-mini API 연동을 통한 위험 사유 텍스트 생성, A/B/C 시나리오 추천 배경 요약, 사후 대처 트리(Fallback Action Plan) 지침 자동 작성"),
    ("프론트엔드 B2B UI 개발", "React 18/19, Zustand, TanStack Table/Query, Recharts 기반 통합 관제, 시뮬레이션 워크벤치, 성과 비교 관제 화면 개발"),
    ("검증 및 최종 보고", "E2E 통합 시나리오 테스트, Oracle Retail 지표 검증, 최종 수행계획서 작성 및 사용자 설명 웹사이트(explainer-site) 구축·배포")
]

# Ensure header
t3.rows[0].cells[0].paragraphs[0].text = " 업무"
t3.rows[0].cells[1].paragraphs[0].text = "업무범위"

# Clear existing rows after header and populate
while len(t3.rows) > 1:
    tr = t3.rows[-1]._tr
    t3._tbl.remove(tr)

for label, val in t3_data:
    row = t3.add_row()
    row.cells[0].paragraphs[0].text = f" {label}"
    row.cells[1].paragraphs[0].text = val

for r in t3.rows:
    for c in r.cells:
        for p in c.paragraphs:
            p.paragraph_format.space_before = Pt(3)
            p.paragraph_format.space_after = Pt(3)
            p.paragraph_format.line_spacing = 1.15
            for run in p.runs:
                run.font.name = "Malgun Gothic"
                run.font.size = Pt(9.5)

# Table 4: Member Roles
t4 = doc.tables[3]
t4_data = [
    ("김준하", "프로젝트 총괄 / 백엔드(Spring Boot, Batch), 파이썬 최적화 엔진, 프론트엔드 B2B UI 개발, DB 스키마 설계"),
    ("김영만", "도메인 멘토 / 현대백화점 직매입 재고 의사결정 정책 수립, Oracle Retail 지표 산정 기준 정의"),
    ("문주성", "오퍼레이션 검증 / 패션 부문 아울렛 이관 권한 및 브랜드 할인 상한 5대 하드 차단 조건 검증")
]

t4.rows[0].cells[0].paragraphs[0].text = " 구분"
t4.rows[0].cells[1].paragraphs[0].text = "역할 및 책임"

while len(t4.rows) > 1:
    tr = t4.rows[-1]._tr
    t4._tbl.remove(tr)

for label, val in t4_data:
    row = t4.add_row()
    row.cells[0].paragraphs[0].text = f" {label}"
    row.cells[1].paragraphs[0].text = val

for r in t4.rows:
    for c in r.cells:
        for p in c.paragraphs:
            p.paragraph_format.space_before = Pt(3)
            p.paragraph_format.space_after = Pt(3)
            for run in p.runs:
                run.font.name = "Malgun Gothic"
                run.font.size = Pt(9.5)

print("Tables updated successfully")
