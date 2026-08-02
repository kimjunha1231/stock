import os
import docx
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import qn, nsdecls

def set_cell_background(cell, color_hex):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)

def set_cell_margins(cell, top=120, bottom=120, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = parse_xml(f'<w:tcMar {nsdecls("w")}><w:top w:w="{top}" w:type="dxa"/><w:bottom w:w="{bottom}" w:type="dxa"/><w:left w:w="{left}" w:type="dxa"/><w:right w:w="{right}" w:type="dxa"/></w:tcMar>')
    tcPr.append(tcMar)

def set_table_borders(table, color="D3D3D3", sz="4", val="single"):
    tblPr = table._tbl.tblPr
    borders = parse_xml(f'''
        <w:tblBorders {nsdecls("w")}>
            <w:top w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>
            <w:bottom w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>
            <w:insideH w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>
            <w:insideV w:val="none"/>
            <w:left w:val="none"/>
            <w:right w:val="none"/>
        </w:tblBorders>
    ''')
    tblPr.append(borders)

def format_cell_text(cell, text, font_name="Malgun Gothic", size_pt=9.5, bold=False, color_rgb=(51,51,51), align=WD_ALIGN_PARAGRAPH.LEFT, space_before=3, space_after=3):
    p = cell.paragraphs[0]
    p.text = ""
    p.alignment = align
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = 1.15
    run = p.add_run(text)
    run.font.name = font_name
    run.font.size = Pt(size_pt)
    run.bold = bold
    run.font.color.rgb = RGBColor(*color_rgb)

def add_paragraph_after_element(doc, text, font_name="Malgun Gothic", size_pt=10, bold=False, space_before=3, space_after=6, color_rgb=(51,51,51), align=WD_ALIGN_PARAGRAPH.LEFT):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = 1.15
    run = p.add_run(text)
    run.font.name = font_name
    run.font.size = Pt(size_pt)
    run.bold = bold
    run.font.color.rgb = RGBColor(*color_rgb)
    return p

# Load original template
doc_path = '/Users/junha/Downloads/프로젝트 수행계획서_양식.docx'
doc = docx.Document(doc_path)

# 1. Update Title Paragraph (P1)
p_title = doc.paragraphs[1]
p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_title.paragraph_format.space_before = Pt(12)
p_title.paragraph_format.space_after = Pt(18)
p_title.text = ""
run_title = p_title.add_run("프 로 젝 트 수 행 계 획 서")
run_title.font.name = "Malgun Gothic"
run_title.font.size = Pt(22)
run_title.bold = True
run_title.font.color.rgb = RGBColor(0, 51, 102)

# 2. Update Table 1 (Key info header)
t1 = doc.tables[0]
set_table_borders(t1, color="B0C4DE")
t1_data = [
    ("팀명", "[2팀] HEYDI Stock (현대백화점 B2B AI 재고 최적화 팀)"),
    ("프로젝트명(주제)", "현대백화점 B2B AI 재고 수익 최적화 플랫폼 (Heydi Stock Pro)"),
    ("멘토", "김영만 선임 (현대백화점그룹 본사 재고전략팀)"),
    ("멘티(교육생)", "김준하 (풀스택 개발자)")
]

for idx, (lbl, val) in enumerate(t1_data):
    format_cell_text(t1.rows[idx].cells[0], lbl, size_pt=9.5, bold=True, color_rgb=(0, 51, 102), align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_background(t1.rows[idx].cells[0], "F0F4F8")
    set_cell_margins(t1.rows[idx].cells[0])
    
    format_cell_text(t1.rows[idx].cells[1], val, size_pt=9.5, bold=False, color_rgb=(51,51,51))
    set_cell_margins(t1.rows[idx].cells[1])

# 3. Update Table 2 (Section I. Project Overview)
t2 = doc.tables[1]
set_table_borders(t2, color="B0C4DE")
t2_data = [
    ("프로젝트명", "현대백화점 B2B AI 재고 수익 최적화 플랫폼 (Heydi Stock Pro)"),
    ("프로젝트 기간", "2026.07.20 ~ 2026.08.15 (4주)"),
    ("목적", "현대백화점 직매입 재고(더현대 서울 2F/3F/B1/1F)의 장기보관 및 유통기한 임박 손실을 새벽 자동 탐지(Spring Batch + Quartz)하고, 수리 최적화 연산기(Python SciPy/PuLP)와 생성형 AI(OpenAI GPT-4o-mini)를 결합하여 Do Nothing Baseline 대비 증분 기여현금이익을 극대화하는 B2B 의사결정 지원 플랫폼 구축"),
    ("키워드", "AI 재고 최적화, 기여현금이익(Contribution Cash Profit), Do Nothing Baseline, 수리 제약 최적화, B2B 오퍼레이션 타워, Oracle Retail RMS"),
    ("핵심기술", "수리 제약 최적화 알고리즘 (SciPy/PuLP), 생성형 AI 판단 근거 연동 (OpenAI GPT-4o-mini API), 대용량 위험재고 자동 탐지 파이프라인 (Spring Batch 5 + Quartz)"),
    ("사용프로그램", "React 18/19, Vite, TypeScript, Zustand, TanStack Table/Query, Recharts, Java 21, Spring Boot 3.3, QueryDSL, Python 3.11+, FastAPI, PostgreSQL 16, Redis, OpenAI API"),
    ("참고문헌", "Character Region Awareness for Text Detection (CVPR 2019), Oracle Retail Merchandising System (RMS) In-Season Inventory Analytics Guide (2024), 현대백화점 지속가능경영보고서 (2024)"),
    ("기대효과", "일일 보관/폐기 손실 사전 회피, Do Nothing 대비 증분 기여현금이익 15% 이상 증대, 5대 하드 차단 제약조건 자동제어로 브랜드 가치 보호, 재고 최적화 의사결정 소요 시간 80% 단축")
]

# Ensure t2 has enough rows
while len(t2.rows) < len(t2_data):
    t2.add_row()

for idx, (lbl, val) in enumerate(t2_data):
    format_cell_text(t2.rows[idx].cells[0], f"  {lbl}", size_pt=9.5, bold=True, color_rgb=(0, 51, 102))
    set_cell_background(t2.rows[idx].cells[0], "F0F4F8")
    set_cell_margins(t2.rows[idx].cells[0])
    
    format_cell_text(t2.rows[idx].cells[1], val, size_pt=9.5, bold=False, color_rgb=(51,51,51))
    set_cell_margins(t2.rows[idx].cells[1])

# 4. Update Table 3 (Section II. Work Breakdown)
t3 = doc.tables[2]
set_table_borders(t3, color="B0C4DE")
t3_data = [
    ("데이터 전처리 및 DB 설계", "더현대 서울 직매입 재고 Master, 보관비/원가/유통기한/시즌 D-Day, 승인이력 PostgreSQL 16 스키마 정규화 설계 및 Mock 데이터 구축"),
    ("백엔드 Core 및 동적 검색 개발", "Spring Boot 3.3, QueryDSL 기반 층/카테고리/D-Day/위험도 다중조건 필터링 API 및 가상 스레드(Virtual Threads) 기반 동시 시뮬레이션 연동"),
    ("새벽 대용량 탐지 배치 개발", "Spring Batch 5 + Quartz Scheduler 기반 매일 02시 전점 재고 일일 보관 손실액 계산 및 위험도 5단계(SAFE~DEAD) 자동 분류 파이프라인 개발"),
    ("수리 최적화 연산 엔진 개발", "Python 3.11 FastAPI 기반 SciPy/PuLP 마진 연산기, 할인율 반응 곡선, 번들 매칭 및 60fps 실시간 슬라이더 시뮬레이션 엔드포인트 구축"),
    ("생성형 AI 판단 근거 연동", "OpenAI GPT-4o-mini API 연동을 통한 위험 사유 텍스트 생성, A/B/C 시나리오 추천 배경 요약, 사후 대처 트리(Fallback Action Plan) 지침 자동 작성"),
    ("프론트엔드 B2B UI 개발", "React 18/19, Zustand, TanStack Table/Query, Recharts 기반 통합 관제, 시뮬레이션 워크벤치, 성과 비교 관제 화면 개발"),
    ("검증 및 최종 보고", "E2E 통합 시나리오 테스트, Oracle Retail 지표 검증, 최종 수행계획서 작성 및 사용자 설명 웹사이트(explainer-site) 구축·배포")
]

format_cell_text(t3.rows[0].cells[0], " 업무", size_pt=9.5, bold=True, color_rgb=(0, 51, 102), align=WD_ALIGN_PARAGRAPH.CENTER)
set_cell_background(t3.rows[0].cells[0], "D9E1F2")
set_cell_margins(t3.rows[0].cells[0])

format_cell_text(t3.rows[0].cells[1], "업무범위", size_pt=9.5, bold=True, color_rgb=(0, 51, 102), align=WD_ALIGN_PARAGRAPH.CENTER)
set_cell_background(t3.rows[0].cells[1], "D9E1F2")
set_cell_margins(t3.rows[0].cells[1])

# Remove old data rows and add t3_data
while len(t3.rows) > 1:
    tr = t3.rows[-1]._tr
    t3._tbl.remove(tr)

for lbl, val in t3_data:
    r = t3.add_row()
    format_cell_text(r.cells[0], f" {lbl}", size_pt=9.5, bold=True, color_rgb=(51,51,51))
    set_cell_margins(r.cells[0])
    format_cell_text(r.cells[1], val, size_pt=9.5, bold=False, color_rgb=(51,51,51))
    set_cell_margins(r.cells[1])

# 5. Update Section II.2 Methodology & Approach (Paragraphs after Table 3)
# We can find paragraph 8 in doc or append right after
p_method = None
for p in doc.paragraphs:
    if 'III. 기반구축' in p.text or 'Ⅲ. 기반구축' in p.text:
        p_method = p
        break

# Let's insert Method text before Section III heading
p_m1 = p_method.insert_paragraph_before()
p_m1.paragraph_format.space_before = Pt(8)
p_m1.paragraph_format.space_after = Pt(4)
r_m1 = p_m1.add_run(" 2. 방법론 및 접근법 ")
r_m1.font.name = "Malgun Gothic"
r_m1.font.size = Pt(11)
r_m1.bold = True
r_m1.font.color.rgb = RGBColor(0, 51, 102)

p_m2 = p_method.insert_paragraph_before()
p_m2.paragraph_format.space_before = Pt(2)
p_m2.paragraph_format.space_after = Pt(3)
p_m2.paragraph_format.line_spacing = 1.15
r_m2 = p_m2.add_run("* 방법론 : 빅데이터 분석 및 Agile 반복 개발 방법론을 채택하여 분석기획, 도메인/데이터 설계, AI/연산 엔진 및 백엔드 구축, 프론트엔드 통합, 검증 및 전개 단계로 순차적·반복적 개발 프로젝트 진행")
r_m2.font.name = "Malgun Gothic"
r_m2.font.size = Pt(9.5)
r_m2.font.color.rgb = RGBColor(51,51,51)

p_m3 = p_method.insert_paragraph_before()
p_m3.paragraph_format.space_before = Pt(2)
p_m3.paragraph_format.space_after = Pt(10)
p_m3.paragraph_format.line_spacing = 1.15
r_m3 = p_m3.add_run("* 접근법 : 하향식 접근법(Top-down)을 채택하여 현황 분석을 통해 도출된 과제(현대백화점 직매입 재고 손실 최소화)에 대하여 프로젝트를 진행하며, 수리 방정식 100% 정밀 연산과 LLM 자연어 지침 생성을 결합한 하이브리드 AI 접근법 적용")
r_m3.font.name = "Malgun Gothic"
r_m3.font.size = Pt(9.5)
r_m3.font.color.rgb = RGBColor(51,51,51)

# 6. Section III: Infrastructure & Architecture (기반구축 내용)
p_sec3_bg = p_method.insert_paragraph_before() # Section III is p_method
p_sec3_bg.paragraph_format.space_before = Pt(4)
p_sec3_bg.paragraph_format.space_after = Pt(3)
r_bg_h = p_sec3_bg.add_run(" 1. 기술개발배경 ")
r_bg_h.font.name = "Malgun Gothic"
r_bg_h.font.size = Pt(11)
r_bg_h.bold = True
r_bg_h.font.color.rgb = RGBColor(0, 51, 102)

p_bg_t1 = p_method.insert_paragraph_before()
p_bg_t1.paragraph_format.space_before = Pt(2)
p_bg_t1.paragraph_format.space_after = Pt(3)
p_bg_t1.paragraph_format.line_spacing = 1.15
r_bg_t1 = p_bg_t1.add_run("* 본 프로젝트는 현대백화점 직매입 재고의 시즌 경과 및 소비기한 임박으로 발생하는 무분별한 덤핑/폐기 손실을 방지하고, 당면한 재고 처리 의사결정 효율성을 극대화하는 것을 일차적인 목표로 함.")
r_bg_t1.font.name = "Malgun Gothic"
r_bg_t1.font.size = Pt(9.5)
r_bg_t1.font.color.rgb = RGBColor(51,51,51)

p_bg_t2 = p_method.insert_paragraph_before()
p_bg_t2.paragraph_format.space_before = Pt(2)
p_bg_t2.paragraph_format.space_after = Pt(3)
p_bg_t2.paragraph_format.line_spacing = 1.15
r_bg_t2 = p_bg_t2.add_run("* 기존 커머스 및 ERP/재고관리 시스템은 단순 일률적 할인이나 단순 폐기에 의존하여 보관비·폐기비·반품비·브랜드 훼손·정상판매 잠식비용을 종합적으로 반영하지 못함.")
r_bg_t2.font.name = "Malgun Gothic"
r_bg_t2.font.size = Pt(9.5)
r_bg_t2.font.color.rgb = RGBColor(51,51,51)

p_bg_t3 = p_method.insert_paragraph_before()
p_bg_t3.paragraph_format.space_before = Pt(2)
p_bg_t3.paragraph_format.space_after = Pt(10)
p_bg_t3.paragraph_format.line_spacing = 1.15
r_bg_t3 = p_bg_t3.add_run("* 따라서 방치 시 기준선(Do Nothing Baseline) 대비 '위험조정 증분 기여현금이익'을 산출하고, 담당자가 직접 조율·검토·승인할 수 있는 B2B AI 오퍼레이션 타워 구축이 필수적임.")
r_bg_t3.font.name = "Malgun Gothic"
r_bg_t3.font.size = Pt(9.5)
r_bg_t3.font.color.rgb = RGBColor(51,51,51)

p_arch_h = p_method.insert_paragraph_before()
p_arch_h.paragraph_format.space_before = Pt(4)
p_arch_h.paragraph_format.space_after = Pt(3)
r_arch_h = p_arch_h.add_run(" 2. 소프트웨어 구성도 ")
r_arch_h.font.name = "Malgun Gothic"
r_arch_h.font.size = Pt(11)
r_arch_h.bold = True
r_arch_h.font.color.rgb = RGBColor(0, 51, 102)

p_arch_t = p_method.insert_paragraph_before()
p_arch_t.paragraph_format.space_before = Pt(2)
p_arch_t.paragraph_format.space_after = Pt(8)
p_arch_t.paragraph_format.line_spacing = 1.15
r_arch_t = p_arch_t.add_run("[ Client: React/Zustand SPA ] <──REST API──> [ Backend: Spring Boot 3.3 / QueryDSL ] <──HTTP──> [ AI Engine: Python FastAPI (SciPy/PuLP) + OpenAI GPT API ]\n                                                          └───────── Database: PostgreSQL 16 & Redis ───────┘\n                                                          └───────── Batch: Spring Batch 5 + Quartz Scheduler ─┘")
r_arch_t.font.name = "Courier New"
r_arch_t.font.size = Pt(8.5)
r_arch_t.font.color.rgb = RGBColor(0, 51, 102)

# 7. Section IV: Implementation System (프로젝트 추진체계)
p_sec4 = None
for p in doc.paragraphs:
    if 'IV. 프로젝트' in p.text or 'Ⅳ. 프로젝트' in p.text:
        p_sec4 = p
        break

# Org chart heading & text
p_org_t = p_sec4.insert_paragraph_before() # Insert before Table 4 / roles heading
# Wait, let's look at P12: ' 1. 프로젝트 수행조직도 '
for p in doc.paragraphs:
    if '1. 프로젝트 수행조직도' in p.text:
        p_org_p = p.insert_paragraph_before()
        p_org_p.paragraph_format.space_before = Pt(4)
        p_org_p.paragraph_format.space_after = Pt(8)
        r_org = p_org_p.add_run("                    [ 프로젝트 수행 총괄 (김준하) ]\n                                   │\n       ┌───────────────────────────┼───────────────────────────┐\n       ▼                           ▼                           ▼\n[ 풀스택 & AI 개발 ]     [ 도메인 & 전략 멘토 ]     [ 오퍼레이션 검증 ]\n (김준하 / 개발총괄)      (김영만 / 수석 MD)        (문주성 / 책임 MD)")
        r_org.font.name = "Courier New"
        r_org.font.size = Pt(8.5)
        r_org.font.color.rgb = RGBColor(0, 51, 102)
        break

# Update Table 4 (Section IV.2 Member Roles)
t4 = doc.tables[3]
set_table_borders(t4, color="B0C4DE")
t4_data = [
    ("김준하", "프로젝트 총괄 / 백엔드(Spring Boot 3.3, Spring Batch), 파이썬 수리 최적화 엔진, 프론트엔드 B2B UI 개발, DB 스키마 설계 및 통합 배포"),
    ("김영만", "도메인 멘토 / 현대백화점 직매입 재고 의사결정 정책 수립, Oracle Retail RMS 분석 지표 정의, 증분 기여현금이익 가드레일 검토"),
    ("문주성", "오퍼레이션 검증 / 더현대 서울 패션 부문 아울렛 이관 권한 및 브랜드 할인 상한 5대 하드 차단 제약조건 검증")
]

format_cell_text(t4.rows[0].cells[0], " 구분", size_pt=9.5, bold=True, color_rgb=(0, 51, 102), align=WD_ALIGN_PARAGRAPH.CENTER)
set_cell_background(t4.rows[0].cells[0], "D9E1F2")
set_cell_margins(t4.rows[0].cells[0])

format_cell_text(t4.rows[0].cells[1], "역할 및 책임", size_pt=9.5, bold=True, color_rgb=(0, 51, 102), align=WD_ALIGN_PARAGRAPH.CENTER)
set_cell_background(t4.rows[0].cells[1], "D9E1F2")
set_cell_margins(t4.rows[0].cells[1])

while len(t4.rows) > 1:
    tr = t4.rows[-1]._tr
    t4._tbl.remove(tr)

for lbl, val in t4_data:
    r = t4.add_row()
    format_cell_text(r.cells[0], f" {lbl}", size_pt=9.5, bold=True, color_rgb=(51,51,51), align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_margins(r.cells[0])
    format_cell_text(r.cells[1], val, size_pt=9.5, bold=False, color_rgb=(51,51,51))
    set_cell_margins(r.cells[1])

# 8. Section V: Management Process (관리 프로세스 계획)
# We need to insert 1. 세부일정 추진 계획 table and 2. 단계별 산출물 table after Section V heading!
p_sec5 = None
for p in doc.paragraphs:
    if 'V. 관리' in p.text or 'Ⅴ. 관리' in p.text:
        p_sec5 = p
        break

p_sch_h = p_sec5.insert_paragraph_before()
p_sch_h.paragraph_format.space_before = Pt(6)
p_sch_h.paragraph_format.space_after = Pt(4)
r_sch_h = p_sch_h.add_run(" 1. 세부일정 추진 계획 ")
r_sch_h.font.name = "Malgun Gothic"
r_sch_h.font.size = Pt(11)
r_sch_h.bold = True
r_sch_h.font.color.rgb = RGBColor(0, 51, 102)

# Schedule Table
t_sch = doc.add_table(rows=7, cols=5)
set_table_borders(t_sch, color="B0C4DE")

sch_headers = ["세부개발내용", "1주 (7/20~7/26)", "2주 (7/27~8/02)", "3주 (8/03~8/09)", "4주 (8/10~8/15)"]
for c_idx, text in enumerate(sch_headers):
    format_cell_text(t_sch.rows[0].cells[c_idx], text, size_pt=9.0, bold=True, color_rgb=(0, 51, 102), align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_background(t_sch.rows[0].cells[c_idx], "D9E1F2")
    set_cell_margins(t_sch.rows[0].cells[c_idx])

sch_rows = [
    ("요구사항 정의 & DB 스키마 설계", "■■■■■", "", "", ""),
    ("백엔드 API & QueryDSL 동적검색 개발", "", "■■■■■", "", ""),
    ("Python 수리 최적화 마진 연산기 개발", "", "■■■■■", "", ""),
    ("Spring Batch 새벽 탐지 & OpenAI API 연동", "", "", "■■■■■", ""),
    ("프론트엔드 B2B UI & 시뮬레이션 워크벤치", "", "", "■■■■■", ""),
    ("E2E 시나리오 테스트 & 산출물/웹사이트 구축", "", "", "", "■■■■■")
]

for r_idx, r_data in enumerate(sch_rows, start=1):
    for c_idx, val in enumerate(r_data):
        align = WD_ALIGN_PARAGRAPH.LEFT if c_idx == 0 else WD_ALIGN_PARAGRAPH.CENTER
        color = (0, 102, 204) if "■" in val else (51, 51, 51)
        bold = True if "■" in val else False
        format_cell_text(t_sch.rows[r_idx].cells[c_idx], val, size_pt=8.5, bold=bold, color_rgb=color, align=align)
        set_cell_margins(t_sch.rows[r_idx].cells[c_idx])

p_out_h = p_sec5.insert_paragraph_before()
p_out_h.paragraph_format.space_before = Pt(10)
p_out_h.paragraph_format.space_after = Pt(4)
r_out_h = p_out_h.add_run(" 2. 단계별 산출물 ")
r_out_h.font.name = "Malgun Gothic"
r_out_h.font.size = Pt(11)
r_out_h.bold = True
r_out_h.font.color.rgb = RGBColor(0, 51, 102)

# Deliverables Table
t_out = doc.add_table(rows=7, cols=4)
set_table_borders(t_out, color="B0C4DE")

out_headers = ["단계", "산출물", "완료일", "비고"]
for c_idx, text in enumerate(out_headers):
    format_cell_text(t_out.rows[0].cells[c_idx], text, size_pt=9.0, bold=True, color_rgb=(0, 51, 102), align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_background(t_out.rows[0].cells[c_idx], "D9E1F2")
    set_cell_margins(t_out.rows[0].cells[c_idx])

out_rows = [
    ("계획수립", "프로젝트 수행계획서", "07/22", "본 문서"),
    ("분석", "요구사항정의서 & 의사결정 정책서", "07/24", "Do Nothing 기준선 정의"),
    ("설계", "DB 스키마 설계서 & AI 최적화 수식 명세서", "07/27", "PostgreSQL / PuLP 연산"),
    ("개발", "백엔드 API & Python AI 엔진 & React B2B 프론트엔드", "08/08", "GitHub 소스코드 반영"),
    ("테스트", "E2E 시나리오 테스트결과서 & 사용설명 웹사이트", "08/12", "explainer-site 배포"),
    ("완료보고", "최종 프로젝트 완료보고서", "08/15", "발표자료 및 완료보고서")
]

for r_idx, r_data in enumerate(out_rows, start=1):
    for c_idx, val in enumerate(r_data):
        align = WD_ALIGN_PARAGRAPH.CENTER if c_idx in [0, 2] else WD_ALIGN_PARAGRAPH.LEFT
        format_cell_text(t_out.rows[r_idx].cells[c_idx], val, size_pt=8.5, bold=False, color_rgb=(51,51,51), align=align)
        set_cell_margins(t_out.rows[r_idx].cells[c_idx])

# Move t_sch and t_out XML elements before Section VI heading
p_sec6 = None
for p in doc.paragraphs:
    if 'VI. 개발' in p.text or 'Ⅵ. 개발' in p.text:
        p_sec6 = p
        break

p_sec6._p.addprevious(p_sch_h._p)
p_sec6._p.addprevious(t_sch._element)
p_sec6._p.addprevious(p_out_h._p)
p_sec6._p.addprevious(t_out._element)

# 9. Section VI: Significance & Importance (개발기술의 의의 및 중요성)
p_sec6_content = p_sec6.insert_paragraph_before()
p_sec6_content.paragraph_format.space_before = Pt(4)
p_sec6_content.paragraph_format.space_after = Pt(3)
r_sig_h = p_sec6_content.add_run("* 개발 기술의 의의")
r_sig_h.font.name = "Malgun Gothic"
r_sig_h.font.size = Pt(10)
r_sig_h.bold = True
r_sig_h.font.color.rgb = RGBColor(0, 51, 102)

p_sig_t1 = p_sec6.insert_paragraph_before()
p_sig_t1.paragraph_format.space_before = Pt(2)
p_sig_t1.paragraph_format.space_after = Pt(3)
p_sig_t1.paragraph_format.line_spacing = 1.15
r_sig_t1 = p_sig_t1.add_run("확률적 LLM의 환각(Hallucination) 위험을 원천 차단하기 위해 '수리 제약 최적화 연산기(100% 수리적 정밀성)'와 '생성형 LLM(자연어 지침 가이드)'을 이원화한 하이브리드 AI 아키텍처를 제시함.")
r_sig_t1.font.name = "Malgun Gothic"
r_sig_t1.font.size = Pt(9.5)
r_sig_t1.font.color.rgb = RGBColor(51,51,51)

p_sig_t2 = p_sec6.insert_paragraph_before()
p_sig_t2.paragraph_format.space_before = Pt(2)
p_sig_t2.paragraph_format.space_after = Pt(8)
p_sig_t2.paragraph_format.line_spacing = 1.15
r_sig_t2 = p_sig_t2.add_run("백엔드(Spring Boot 3.3)와 AI 엔진(Python FastAPI)을 MSA 구조로 독립 분리하여, 수리 최적화 알고리즘 수정 시에도 메인 시스템 재부팅 없이 유연한 확장 및 배포가 가능함.")
r_sig_t2.font.name = "Malgun Gothic"
r_sig_t2.font.size = Pt(9.5)
r_sig_t2.font.color.rgb = RGBColor(51,51,51)

p_imp_h = p_sec6.insert_paragraph_before()
p_imp_h.paragraph_format.space_before = Pt(4)
p_imp_h.paragraph_format.space_after = Pt(3)
r_imp_h = p_imp_h.add_run("* 중요성")
r_imp_h.font.name = "Malgun Gothic"
r_imp_h.font.size = Pt(10)
r_imp_h.bold = True
r_imp_h.font.color.rgb = RGBColor(0, 51, 102)

p_imp_t1 = p_sec6.insert_paragraph_before()
p_imp_t1.paragraph_format.space_before = Pt(2)
p_imp_t1.paragraph_format.space_after = Pt(10)
p_imp_t1.paragraph_format.line_spacing = 1.15
r_imp_t1 = p_imp_t1.add_run("현대백화점 패션/식품 직매입 재고에 Oracle Retail 기준 5대 분석지표(ROS, Sell Thru %, WOS, Lift %, STD/BTA)와 5대 하드 차단 제약조건(식품안전, 아울렛권한, 물류용량, 브랜드상한, AI ROI)을 최초로 정밀 적용하여 실무 현장 도입 가능성을 극대화함.")
r_imp_t1.font.name = "Malgun Gothic"
r_imp_t1.font.size = Pt(9.5)
r_imp_t1.font.color.rgb = RGBColor(51,51,51)

# 10. Section VII: Application & Expected Effects (활용방안/적용분야 및 기대효과)
p_sec7 = None
for p in doc.paragraphs:
    if 'VII' in p.text or 'Ⅶ' in p.text:
        p_sec7 = p
        break

p_app_h = p_sec7.insert_paragraph_before()
p_app_h.paragraph_format.space_before = Pt(4)
p_app_h.paragraph_format.space_after = Pt(3)
r_app_h = p_app_h.add_run("* 활용방안")
r_app_h.font.name = "Malgun Gothic"
r_app_h.font.size = Pt(10)
r_app_h.bold = True
r_app_h.font.color.rgb = RGBColor(0, 51, 102)

p_app_t1 = doc.add_paragraph()
p_app_t1.paragraph_format.space_before = Pt(2)
p_app_t1.paragraph_format.space_after = Pt(3)
p_app_t1.paragraph_format.line_spacing = 1.15
r_app_t1 = p_app_t1.add_run("더현대 서울 4개 층(2F 여성패션, 3F 남성/잡화, B1 식품관, 1F 뷰티/리빙) 직매입 재고 관제 적용을 시작으로 현대백화점 15개 점포, 9개 현대아울렛, 현대식품관/Hmall 등 전사 유통 망으로 확산 가능.")
r_app_t1.font.name = "Malgun Gothic"
r_app_t1.font.size = Pt(9.5)
r_app_t1.font.color.rgb = RGBColor(51,51,51)

p_app_t2 = doc.add_paragraph()
p_app_t2.paragraph_format.space_before = Pt(2)
p_app_t2.paragraph_format.space_after = Pt(8)
p_app_t2.paragraph_format.line_spacing = 1.15
r_app_t2 = p_app_t2.add_run("직매입 수치 연산 체계를 바탕으로 입점 브랜드 특약매입 재고 대상 제안형 판촉 및 프로모션 시뮬레이션 솔루션으로 연계 활용 가능.")
r_app_t2.font.name = "Malgun Gothic"
r_app_t2.font.size = Pt(9.5)
r_app_t2.font.color.rgb = RGBColor(51,51,51)

p_eff_h = doc.add_paragraph()
p_eff_h.paragraph_format.space_before = Pt(4)
p_eff_h.paragraph_format.space_after = Pt(3)
r_eff_h = p_eff_h.add_run("* 기대효과")
r_eff_h.font.name = "Malgun Gothic"
r_eff_h.font.size = Pt(10)
r_eff_h.bold = True
r_eff_h.font.color.rgb = RGBColor(0, 51, 102)

p_eff_t1 = doc.add_paragraph()
p_eff_t1.paragraph_format.space_before = Pt(2)
p_eff_t1.paragraph_format.space_after = Pt(3)
p_eff_t1.paragraph_format.line_spacing = 1.15
r_eff_t1 = p_eff_t1.add_run("악성 재고 보관/폐기 손실액 25% 이상 감축, 아무것도 하지 않는 방치(Do Nothing) 대비 증분 기여현금이익 15% 이상 증대.")
r_eff_t1.font.name = "Malgun Gothic"
r_eff_t1.font.size = Pt(9.5)
r_eff_t1.font.color.rgb = RGBColor(51,51,51)

p_eff_t2 = doc.add_paragraph()
p_eff_t2.paragraph_format.space_before = Pt(2)
p_eff_t2.paragraph_format.space_after = Pt(12)
p_eff_t2.paragraph_format.line_spacing = 1.15
r_eff_t2 = p_eff_t2.add_run("5대 하드 차단 제약조건 자동 검증을 통한 브랜드 가치 보호 및 법규 준수 보장, 재고 최적화 의사결정 소요 시간 80% 단축.")
r_eff_t2.font.name = "Malgun Gothic"
r_eff_t2.font.size = Pt(9.5)
r_eff_t2.font.color.rgb = RGBColor(51,51,51)

# Save document
output_path = '/Users/junha/coding/stock/deliverables/프로젝트 수행계획서_양식.docx'
doc.save(output_path)

print(f"Document successfully written and saved to {output_path}")
